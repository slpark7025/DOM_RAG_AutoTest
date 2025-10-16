"""자동화 수행을 위한 기본 셋팅 내용 포함"""

import sys
import os
import glob

from numpy.ma.testutils import assert_equal
import testrail_setting
import log_setting
import move_menu
import time
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException
import zipfile
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from urllib.parse import urlparse
from docx import Document
from selenium.webdriver.chrome.options import Options
from openpyxl import load_workbook, Workbook
import os
import gspread
from google.oauth2.service_account import Credentials

# 자동화 실행 시 전달 받은 인자 값을 저장
address = "http://" + sys.argv[1] # VPES 서버 정보
run_id = sys.argv[2] # Testrail Run ID 정보
account_info = sys.argv[3] # VPES id정보 입력

# TestRail 정보 및 Message 정보
client = testrail_setting.APIClient('http://testrail.suresoft.intra/testrail')
passMsg = 'Test Successful'

# 테스트 및 테스트 결과 입력을 위한 정보
addressLogin = address + "/vpes/login"
tr_ass = 4
tr_version = "Auto"
client.user = "sun@suresofttech.com"
client.password = "Suresoft1!"

# VPES 서버 아이디, 패스워드 정보
if account_info == "sqa-vpes":
    usr = "admin"
    pwd = "suresoft"
else:
    usr = account_info
    pwd = "qwer1234!!"

def setup(report_type="report"):
    default_path()
    directory = ""
    current_directory = os.getcwd()
    if report_type == "report":
        directory = os.getcwd() + "\\" + 'report'
    user_profile_directory = os.path.join(current_directory, "Profile") # 시용자 프로필 지정
    options = webdriver.ChromeOptions()
    # 창 숨기는 옵션 - 원하면 주석 해제
    #options.add_argument('--headless')
    options.add_argument(f"user-data-dir={user_profile_directory}") # 사용자 프로필 지정
    options.add_experimental_option("prefs", {
        "download.prompt_for_download": False,
        "download.default_directory": directory,
    })

    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    driver.maximize_window()
    driver.implicitly_wait(30)
    return driver


def default_path():
    os.chdir(os.path.dirname(__file__))

def update_testrail_test_run():
    # project id 및 suite id를 통하여 automation_type 필드가 'Selenium'인 TC ID 추출
    count = 0
    get_case = client.send_get('get_cases/1&suite_id=1')
    case_list = []
    for auto_type in get_case:
        if auto_type['custom_automation_type'] == 14:
            case_list.append(auto_type['id'])
            count += 1

    # run id를 통해 plan id 추출
    a = client.send_get('get_run/%s' % run_id)
    get_plan_id = a['plan_id']

    # plan id로 plan 정보를 추출하고 그 정보로 entry_id를 추출
    entry_id = ''
    get_plan = client.send_get('get_plan/%s' % get_plan_id)
    plan_entries = get_plan['entries']
    for entries_list in plan_entries:
        for entries_dict in entries_list['runs']:
            if entries_dict['id'] == int(run_id):
                entry_id = entries_dict['entry_id']
            else:
                pass

    # case_list 리스트의 목록을 이용하여 TestRun 업데이트
    client.send_post('update_plan_entry/%s/%s' % (get_plan_id, entry_id),
                     {'case_ids': case_list})

def get_result(case):
    # Get the test result
    if hasattr(case, "_outcome"):  # Python 3.4+
        result = case.defaultTestResult()
        case._feedErrorsToResult(result, case._outcome.errors)
    else:  # Python 3.2 - 3.3 or older
        result = getattr(case, "_outcomeForDoCleanups", case._resultForDoCleanups)

    # Define the Excel file path
    excel_path = "./log/test_results.xlsx"
    os.makedirs(os.path.dirname(excel_path), exist_ok=True)
    print("엑셀 저장 경로:", os.path.abspath(excel_path))

    # Open or create the workbook
    if os.path.exists(excel_path):
        wb = load_workbook(excel_path)
    else:
        wb = Workbook()
        ws = wb.active
        ws.append(["Case ID", "Result", "Details"])  # Add headers
    ws = wb.active

    # Extract the case ID
    test_name = case.id().split('.')[1]  # 'C282502' 부분 추출

    # Determine pass/fail status
    error_msg = None
    for test_case, exc_info in result.errors + result.failures:
        if test_case is case:
            error_msg = str(exc_info)
            break

    # Find the matching row and update the result
    updated = False
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=1):
        if row[0].value == test_name:
            ws.cell(row=row[0].row, column=2, value="FAIL" if error_msg else "PASS")
            ws.cell(row=row[0].row, column=3, value=error_msg if error_msg else "PASS")
            updated = True
            break

    # If no matching row was found, add a new one
    if not updated:
        ws.append([test_name, "FAIL" if error_msg else "PASS", error_msg if error_msg else "PASS"])

    # Save the workbook
    wb.save(excel_path)
    print(f"Test result for '{test_name}' saved successfully.")

    return result

def upload_result(case, case_id, result):
    # 구글 시트 인증
    json_path = os.path.join(os.path.dirname(__file__), "JSONFILE/vpes-421022-f74d5a7b2e05.json")
    scopes = ["https://www.googleapis.com/auth/spreadsheets"]
    credentials = Credentials.from_service_account_file(json_path, scopes=scopes)
    gc = gspread.authorize(credentials)

    # 시트 열기
    sheet_key = "1GDDlTp2x4BWnLCeNhCvJenU_kfwD-PeYr2hrksCdh5w"  # 공유된 구글 시트 키
    worksheet = gc.open_by_key(sheet_key).worksheet("vpes")  # 탭 이름 주의 (예: 'vpes')

    # 결과 처리
    error_msg = None
    for test_case, exc_info in result.errors + result.failures:
        if test_case is case:
            error_msg = str(exc_info)
            break
    status = "FAIL" if error_msg else "PASS"
    details = error_msg if error_msg else "PASS"

    # case_id가 이미 존재하는지 확인
    try:
        col_values = worksheet.col_values(1)  # 1열(Case ID) 가져오기
        found = False
        for idx, val in enumerate(col_values):
            if val == case_id:
                worksheet.update_cell(idx + 1, 2, status)   # B열: 결과
                worksheet.update_cell(idx + 1, 3, details)  # C열: 상세
                found = True
                print(f"✅ 기존 행 덮어쓰기 완료: row={idx + 1}, case_id={case_id}")
                break

        if not found:
            worksheet.append_row([case_id, status, details])
            print(f"➕ 새 행 추가 완료: case_id={case_id}")
    except Exception as e:
        print("❌ 구글 시트 업데이트 중 오류:", e)
# def upload_result(case, case_id, result):
#     def list2reason(exc_list):
#         if exc_list and exc_list[-1][0] is case:
#             return exc_list[-1][1]
#
#     error = list2reason(result.errors)
#     failure = list2reason(result.failures)
#     if not error and not failure:
#         send_testrail_result(1, case_id, comment=True)
#     else:
#         cmt = failure if error is None else error
#         send_testrail_result(5, case_id, comment=cmt)
#
# def reset_testrail_result():
#     count = 0
#     dir_list = search_test_suite_list()
#     # C* 파일 찾아서 TestRun 초기화
#     for suite_path in dir_list:
#         os.chdir(suite_path)
#         find_ext = 'C*.py'  # 찾고 싶은 파일
#         # 확장자 및 식별자(C) 제거
#         for a in glob.glob(os.path.join(suite_path, find_ext)):
#             b = os.path.basename(a)[:-3]
#             case_id = os.path.basename(b)[1:]
#
#             client.send_post(
#                 'add_result_for_case/%s/%s' % (run_id, case_id),
#                 {'status_id': 8, 'version': tr_version, 'assignedto_id': tr_ass})
#             log_setting.logger.info(b + " 테스트 초기화 완료!")
#             count += 1
#     log_setting.logger.info("테스트 케이스 결과 " + str(count) + " 건 초기화 완료!")


def send_testrail_result(status_id, case_id, comment):
    tr_case_id = case_id[1:]
    cmt = str(comment)
    # Test Rail 결과 입력
    if status_id == 1:
        log_setting.logger.info('\nRun ID : %s\nTest Case ID: %s\nMessage : %s\n' % (run_id, case_id, passMsg))
        client.send_post(
            'add_result_for_case/%s/%s' % (run_id, tr_case_id),
            {'status_id': status_id, 'version': tr_version, 'assignedto_id': tr_ass})

    elif status_id == 5:
        log_setting.logger.info('\nRun ID : %s\nTest Case ID: %s\nMessage : %s\n' % (run_id, case_id, cmt))
        client.send_post(
            'add_result_for_case/%s/%s' % (run_id, tr_case_id),
            {'status_id': status_id, 'comment': cmt, 'version': tr_version, 'assignedto_id': tr_ass})

def search_test_suite_list():
    default_path()
    os.chdir('./test_suite')
    dir_list = []
    target_dir = os.path.normpath(os.getcwd())
    for fname in os.listdir(target_dir):
        if os.path.isdir(fname):
            if fname in '__pycache__':
                pass
            else:
                full_dir = os.path.join(target_dir, fname)
                dir_list.append(full_dir)
        else:
            pass

    return dir_list

# 로그인
def login(driver):
    # VPES 서버 열기 및 로그인
    driver.get(addressLogin)
    driver.implicitly_wait(30)

    driver.find_element(By.ID, "username").send_keys(usr)
    driver.find_element(By.ID, "pwd").send_keys(pwd)
    driver.find_element(By.XPATH, '//button[contains(text(), "로그인")]').click()
    time.sleep(2)
    log_setting.logger.info("The account is available.")

# 프로젝트 삭제 - 추후 정리
def delete_project(driver, case_id):
    move_menu.move_projectSetting(driver, case_id)
    driver.find_element(By.ID,'dropdownMenuButton').click()
    driver.find_element(By.XPATH, "//*/text()[normalize-space(.)='프로젝트 삭제']/parent::*").click()
    time.sleep(2)
    driver.find_element(By.ID,'projectKeyForDelete').send_keys(case_id)
    driver.find_element(By.XPATH, '//*[@id="projectDeleteModal"]/div/div/div[3]/button[2]').click()

# 회차 종료
def close_episode(driver, case_id):
    move_menu.move_progress(driver, case_id)
    driver.find_element(By.XPATH,"//*[@id='projectDetailTransition-HeaderVue']//button[2]").click()
    time.sleep(1)
    driver.find_element(By.ID, "deleteY").click()
    time.sleep(3)


def create_project(driver, case_id, range, scm_type, file_name = None, input_text = None):
    # input_text 변경시 Ex) default_setting.create_project(driver,case_id, "CSCI", "GIT", None, case_id) 형식 사용
    if input_text == None:
        input_text = "test" # 모든 입력값은 "Test"로 통일
    git_url = "https://github.com/YunNaRi/VPES_Autumation.git"
    svn_url = "https://10.10.111.235/svn/VPES_Auto/"

    move_menu.move_project_registration(driver)

    # 프로젝트 명 입력 후 중복 확인
    projectKey_check = driver.find_element(By.ID, "projectRegister-projectKey")
    projectKey_check.send_keys(case_id)
    driver.find_element(By.ID, "projectRegister-duplicationBtn").click()
    time.sleep(1)
    # 프로젝트 명 중복 확인 결과 확인
    project_img = projectKey_check.get_attribute('class')
    if 'is-valid' not in project_img:
        delete_project(driver, case_id) #중복 시 프로젝트 삭제
        create_project(driver,case_id, range, scm_type, file_name, input_text) # 다시 생성
        return True

    time.sleep(1)
    driver.find_element(By.ID, "projectRegister-systemName").send_keys(input_text)
    # 기본값이 CSCI이므로 CSCI에 값을 입력
    driver.find_element(By.ID, "projectRegister-csciName").send_keys(input_text)

    if range == "CSC":
        driver.find_element(By.XPATH,'//*[@id="projectRegister-requiredOpt"]/div[1]/div[3]/div[1]/div/div/div[2]').click() # Combobox선택
        time.sleep(1)
        driver.find_element(By.XPATH,"//div[1]/div[3]/div[1]/div/div/div[3]/ul/li[2]/span").click()
        time.sleep(1)
        driver.find_element(By.ID, "projectRegister-cscName").send_keys(input_text)
    elif range == "CSU":
        driver.find_element(By.XPATH,'//*[@id="projectRegister-requiredOpt"]/div[1]/div[3]/div[1]/div/div/div[2]').click() # Combobox선택
        time.sleep(1)
        driver.find_element(By.XPATH,"//div[1]/div[3]/div[1]/div/div/div[3]/ul/li[3]/span").click()
        time.sleep(1)
        driver.find_element(By.ID, "projectRegister-cscName").send_keys(input_text)
        driver.find_element(By.ID, "projectRegister-csuName").send_keys(input_text)

    time.sleep(3)
    driver.find_element(By.ID, "projectRegister-registerProjectBtn").click() # 프로젝트 생성
    time.sleep(1.5)
    driver.find_element(By.ID, "projectRegister-successBtn").click()
    time.sleep(1.5)

    setting_SCM(driver, case_id, scm_type, file_name, input_text)

def setting_SCM(driver, case_id, scm_type, file_name = None, input_text = None):
    # input_text 변경시 Ex) default_setting.create_project(driver,case_id, "CSCI", "GIT", None, case_id) 형식 사용
    if input_text == None:
        input_text = "test" # 모든 입력값은 "Test"로 통일
    git_url = "https://github.com/YunNaRi/VPES_Autumation.git"
    svn_url = "https://10.10.111.235/svn/VPES_Auto/"

    move_menu.move_project_scm_setting(driver, case_id)

    scm_url_check = driver.find_element(By.ID, "projectRegister-scmUrl")

    if scm_type == "GIT":
        scm_url_check.send_keys(git_url)
        driver.find_element(By.ID, "projectRegister-scmID").send_keys(input_text)
        driver.find_element(By.ID, "projectRegister-scmPW").send_keys(input_text)
    elif scm_type == "SVN":
        driver.find_element(By.XPATH,'//*[@id="configurationManagementModal"]/div/div/div[2]/div[2]/div[2]/div[2]/div[1]/div/div/div[2]').click() # Combobox선택
        time.sleep(1)
        driver.find_element(By.XPATH,"//*[@id='configurationManagementModal']/div/div/div[2]/div[2]/div[2]/div[2]/div[1]/div//ul//span[contains(text(),'SVN')]").click()
        if file_name == None:
            file_name = ''
        scm_url_check.send_keys(svn_url + file_name)
        driver.find_element(By.ID, "projectRegister-scmID").send_keys(input_text)
        driver.find_element(By.ID, "projectRegister-scmPW").send_keys(input_text)

    # 인증 버튼 선택
    driver.find_element(By.ID, "projectRegister-scmAuthBtn").click()
    time.sleep(1)
    # SCM 인증 결과 확인
    scm_img = scm_url_check.get_attribute('class')
    if 'is-valid' not in scm_img:
        return False

    time.sleep(3)

    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[2]").click()
    time.sleep(1)
    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[1]").click()

def setting_SCM_detail(driver, case_id, scm_type, scm_url, file_name = None, input_text = None):
    # input_text 변경시 Ex) default_setting.create_project(driver,case_id, "CSCI", "GIT", None, case_id) 형식 사용
    if input_text == None:
        input_text = "test" # 모든 입력값은 "Test"로 통일
    git_url = scm_url
    svn_url = scm_url

    move_menu.move_project_scm_setting(driver, case_id)

    scm_url_check = driver.find_element(By.ID, "projectRegister-scmUrl")

    if scm_type == "GIT":
        scm_url_check.clear()
        scm_url_check.send_keys(git_url)
        driver.find_element(By.ID, "projectRegister-scmID").clear()
        driver.find_element(By.ID, "projectRegister-scmID").send_keys(input_text)
        driver.find_element(By.ID, "projectRegister-scmPW").clear()
        driver.find_element(By.ID, "projectRegister-scmPW").send_keys(input_text)
    elif scm_type == "SVN":
        driver.find_element(By.XPATH,'//*[@id="configurationManagementModal"]/div/div/div[2]/div[2]/div[2]/div[2]/div[1]/div/div/div[2]').click() # Combobox선택
        time.sleep(1)
        driver.find_element(By.XPATH,"//*[@id='configurationManagementModal']/div/div/div[2]/div[2]/div[2]/div[2]/div[1]/div//ul//span[contains(text(),'SVN')]").click()
        if file_name == None:
            file_name = ''
        scm_url_check.clear()
        scm_url_check.send_keys(svn_url + file_name)
        driver.find_element(By.ID, "projectRegister-scmID").clear()
        driver.find_element(By.ID, "projectRegister-scmID").send_keys(input_text)
        driver.find_element(By.ID, "projectRegister-scmPW").clear()
        driver.find_element(By.ID, "projectRegister-scmPW").send_keys(input_text)

    # 인증 버튼 선택
    driver.find_element(By.ID, "projectRegister-scmAuthBtn").click()
    time.sleep(1)
    # SCM 인증 결과 확인
    scm_img = scm_url_check.get_attribute('class')
    if 'is-valid' not in scm_img:
        return False

    time.sleep(3)

    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[2]").click()
    time.sleep(1)
    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[1]").click()


def create_project_dir(driver, case_id, range, directory_name, type =None, input_text = None):
    # input_text 변경시 Ex) default_setting.create_project_dir(driver, case_id, '', "miro_vs2015", None, case_id) 형식 사용
    if input_text == None:
        input_text = "test" # 모든 입력값은 "Test"로 통일

    if type == '직접 입력':
        text = '직접 입력'

    elif type == None:
        text = 'DEFAULT'

    else:
        text = type

    move_menu.move_project_registration(driver)

    # 프로젝트 명 입력 후 중복 확인
    projectKey_check = driver.find_element(By.ID, "projectRegister-projectKey")
    projectKey_check.send_keys(case_id)
    driver.find_element(By.ID, "projectRegister-duplicationBtn").click()
    time.sleep(1)
    # 프로젝트 명 중복 확인 결과 확인
    project_img = projectKey_check.get_attribute('class')
    if 'is-valid' not in project_img:
        delete_project(driver, case_id)  # 중복 시 프로젝트 삭제
        create_project_dir(driver, case_id, range, directory_name, type, input_text) # 다시 생성
        return True

    time.sleep(1)
    driver.find_element(By.ID, "projectRegister-systemName").send_keys(input_text)
    # 기본값이 CSCI이므로 CSCI에 값을 입력
    driver.find_element(By.ID, "projectRegister-csciName").send_keys(input_text)

    if range == "CSC":
        driver.find_element(By.XPATH,'//*[@id="projectRegister-requiredOpt"]/div[1]/div[3]/div[1]/div/div/div[2]').click() # Combobox선택
        time.sleep(1)
        driver.find_element(By.XPATH,"//div[1]/div[3]/div[1]/div/div/div[3]/ul/li[2]/span").click()
        time.sleep(1)
        driver.find_element(By.ID, "projectRegister-cscName").send_keys(input_text)
    elif range == "CSU":
        driver.find_element(By.XPATH,'//*[@id="projectRegister-requiredOpt"]/div[1]/div[3]/div[1]/div/div/div[2]').click() # Combobox선택
        time.sleep(1)
        driver.find_element(By.XPATH,"//div[1]/div[3]/div[1]/div/div/div[3]/ul/li[3]/span").click()
        time.sleep(1)
        driver.find_element(By.ID, "projectRegister-cscName").send_keys(input_text)
        driver.find_element(By.ID, "projectRegister-csuName").send_keys(input_text)

    time.sleep(2)
    driver.find_element(By.ID, "projectRegister-registerProjectBtn").click() # 프로젝트 생성
    time.sleep(0.5)
    driver.find_element(By.ID, "projectRegister-successBtn").click()
    time.sleep(1)

    setting_SCM_DIR(driver, case_id, directory_name, text)


def setting_SCM_DIR(driver, case_id, directory_name, type =None):

    move_menu.move_project_scm_setting(driver, case_id)

    if type == '직접 입력':
        text = '직접 입력'

    elif type == None:
        text = 'DEFAULT'

    else:
        text = type

    # SCM 설정
    scm_url_check = driver.find_element(By.ID, "projectRegister-scmUrl")
    driver.find_element(By.XPATH, '//*[@id="configurationManagementModal"]/div/div/div[2]/div[2]/div[2]/div[2]/div[1]/div/div/div[2]').click()  # Combobox선택
    time.sleep(1)
    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[2]/div[2]/div[2]/div[2]/div[1]/div//ul//span[contains(text(),'DIRECTORY')]").click()

    # dir 선택
    driver.find_element(By.XPATH, '//*[@id="configurationManagementModal"]/div/div/div[2]/div[2]/div[2]/div[2]/div[2]/div/div/div').click()
    time.sleep(1)
    driver.find_element(By.XPATH, f"//span/span[text()='{text}']").click()
    scm_url_check.send_keys(directory_name)

    time.sleep(1)

    # 인증 버튼 선택
    driver.find_element(By.ID, "projectRegister-scmAuthBtn").click()
    time.sleep(1)
    # SCM 인증 결과 확인
    scm_img = scm_url_check.get_attribute('class')
    if 'is-valid' not in scm_img:
        return False

    time.sleep(3)
    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[2]").click()
    time.sleep(1)
    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[1]").click()
    time.sleep(1)

def setting_multi_SCM_dir(driver, case_id, directory_name, type =None):
    # input_text 변경시 Ex) default_setting.create_project(driver,case_id, "CSCI", "GIT", None, case_id) 형식 사용
    if type == '직접 입력':
        text = '직접 입력'

    elif type == None:
        text = 'DEFAULT'

    else:
        text = type

    move_menu.move_project_scm_setting(driver, case_id)

    # SCM 설정 추가 클릭 - SCM 설정 추가
    driver.find_element(By.XPATH, '//*[@id="configurationManagementModal"]/div/div/div[2]/div[1]/div/div/div/div/div[1]').click()  # Combobox선택
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="configurationManagementModal"]/div/div/div[2]/div[1]/div/div/div/div/div[3]/ul/li[3]/span/span').click()

    multi_scm_url_check = driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[2]/div[3]//*[@id='projectRegister-scmUrl']")

    # dir 선택
    driver.find_element(By.XPATH, '//*[@id="configurationManagementModal"]/div/div/div[2]/div[3]/div[2]/div[2]/div[2]/div/div/div/div[1]').click()
    time.sleep(1)
    driver.find_element(By.XPATH, f"//span[text()='{text}']").click()
    multi_scm_url_check.send_keys(directory_name)

    # 인증 버튼 선택
    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[2]/div[3]//*[@id='projectRegister-scmAuthBtn']").click()
    time.sleep(1)
    # SCM 인증 결과 확인
    scm_img = multi_scm_url_check.get_attribute('class')
    if 'is-valid' not in scm_img:
        return False

    time.sleep(3)

    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[2]").click()
    time.sleep(1)
    driver.find_element(By.XPATH, "//*[@id='configurationManagementModal']/div/div/div[3]/button[1]").click()

def upload_project(driver, case_id, type, file_name):
    tool_setting(driver, case_id, type)
    # 프로젝트 개요 페이지 이동
    move_menu.move_progress(driver, case_id)

    # 신뢰성 시험 결과 클릭
    driver.find_element(By.XPATH,'//div[2]/div[2]/div[2]/button').click()
    time.sleep(1)

    # 업로드 클릭
    driver.find_element(By.XPATH,"//a[contains(text(),'업로드')]").click()

    # 업로드할 도구 클릭
    driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='STATIC']/parent:: *").click()
    if type.lower() == 'static':
        driver.find_element(By.XPATH,'//div[14]/div/div/div/div[2]/div/div/div/div[3]/ul/li/span').click()
        directory = 'STATIC'

    elif type == 'Code Inspector' or type == 'CI':
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='Code Inspector']/parent::*").click()
        directory = 'CI'

    elif type == 'SNIPER' or type == 'SN':
        driver.find_element(By.XPATH,'//div[14]/div/div/div/div[2]/div/div/div/div[3]/ul/li[3]/span').click()
        directory = 'SN'

    elif type == 'Controller Tester' or type == 'CT':
        driver.find_element(By.XPATH,'//*[@id="projectDetailOverView-uploadDownloadModal"]/div/div/div[2]/div[1]/div/div/div[3]/ul/li[4]').click()
        directory = 'CT'

    else:
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='COVER EE/SE']/parent::*").click()
        directory = 'COVER'

    # 해당 도구에 맞는 경로에서 xml 파일 가져오기
    script_directory = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_directory, 'xml', directory, file_name)

    absolute_path = file_path + ".xml"
    driver.find_element(By.ID, "projectDetailOverView-resultFile").send_keys(absolute_path)

    # 업로드 클릭
    driver.find_element(By.ID, 'projectDetailOverView-uploadBtn').click()

    while True:
        try:
            # uploadResultList 엘리먼트 찾기
            upload_result_list = driver.find_element(By.ID, 'uploadResultList')

            # uploadResultList 엘리먼트에 text 값이 있으면 루프 종료
            if upload_result_list.text.strip():
                break

            time.sleep(1)

        # 못찾으면 false
        except NoSuchElementException:
            False


    time.sleep(5)

    # 신뢰성 시험 결과 창 닫기
    driver.find_element(By.XPATH,'//div[14]/div/div/div/div/button/img').click()

    # 프로젝트 상태가 진행중일 경우에 while문 종료.
    progress_status = driver.find_element(By.XPATH,'//*[@id="projectSideBar"]/div/span[1]')
    while True:
        if progress_status.text == "진행중":
            break

        time.sleep(1)

# 계정 삭제
def account_delete (driver, case_id):
    move_menu.move_globalSetting(driver)
    # account_id 검색
    driver.find_element(By.XPATH,"//div[2]/div[2]/label/input").send_keys(case_id)
    driver.find_element(By.XPATH,"//div[2]/div[2]/label/input").send_keys(Keys.ENTER)
    time.sleep(0.5)
    # 검색 결과 확인 후 존재 시 account_id 계정 삭제
    account_check = driver.find_element(By.ID, "globalSettingUserContent")
    if "데이터가 존재하지 않습니다." not in account_check.text:
        if "검색할 데이터가 없습니다." not in account_check.text:
            # 검색 결과가 여러개 나올 경우
            account_list = driver.find_elements(By.XPATH, "//img[contains(@onclick, "+case_id+")]")
            for a in account_list:
                if a.get_attribute('id') == 'userDelete':
                    a.click()  # 아이디값이 account_id인 결과값의 삭제 버튼 선택
                    time.sleep(0.5)
                    driver.find_element(By.ID, "deleteY").click()
                    time.sleep(0.5)
            try:
                toast_element = driver.find_element(By.ID, "toastBottomCenter_COMMON_ERROR0").get_attribute("class")
                # show 클래스가 포함된 경우
                if "show" in toast_element:
                    user_login(driver, case_id)
                    delete_project(driver, case_id)
                    login(driver)
                    account_delete(driver, case_id)
            except NoSuchElementException:
                print("요소 없음")


# 파일 삭제
def file_delete(file_path):
    try:
        os.remove(file_path)
        print(f"{file_path} 파일이 삭제되었습니다.")
    except FileNotFoundError:
        print(f"{file_path} 파일이 이미 삭제되었거나 존재하지 않습니다.")
    except Exception as e:
        print(f"파일 삭제 중 오류가 발생했습니다: {str(e)}")

def download_project(driver, case_id, type):
    tool_setting(driver, case_id, type)
    # 프로젝트 개요 페이지 이동
    move_menu.move_progress(driver, case_id)

    # 신뢰성 시험 결과 클릭
    driver.find_element(By.XPATH,'//div[2]/div[2]/div[2]/button').click()
    time.sleep(1)

    # 다운로드 클릭
    driver.find_element(By.XPATH,"//a[contains(text(),'다운로드')]").click()

    # 다운로드할 도구 클릭
    driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='STATIC']/parent:: *").click()
    if type.lower() == 'static':
        driver.find_element(By.XPATH,'//div[14]/div/div/div/div[2]/div/div/div/div[3]/ul/li[2]/span').click()

    elif type == 'Code Inspector' or type == 'CI':
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='Code Inspector']/parent::*").click()
        directory = 'CI'

    elif type == 'SNIPER' or type == 'SN':
        driver.find_element(By.XPATH,'//div[14]/div/div/div/div[2]/div/div/div/div[3]/ul/li[4]/span').click()

    elif type == 'Controller Tester' or type == 'CT':
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='CT']/parent::*").click()
        directory = 'CT'

    elif type == '전체 다운로드':
        driver.find_element(By.XPATH,'//div[14]/div/div/div/div[2]/div/div/div/div[3]/ul/li/span').click()

    else:
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='COVER EE/SE']/parent::*").click()

    # 다운로드 클릭
    driver.find_element(By.ID, 'projectDetailOverView-downloadBtn').click()

    time.sleep(5)

    # 신뢰성 시험 결과 창 닫기
    driver.find_element(By.XPATH,'//*[@id="projectDetailOverView-uploadDownloadModal"]/div/div/div[1]/button/img').click()


def tool_setting(driver, case_id, type):
    move_menu.move_projectSetting(driver, case_id)
    driver.find_element(By.ID, 'projectRegister-additionalOpt-tab').click()

    driver.find_element(By.XPATH,"(.//*[normalize-space(text()) and normalize-space(.)='코딩 규칙'])[1]/following::span[1]").click()
    time.sleep(0.5)

    if type.lower() == 'static':
        driver.find_element(By.XPATH,'//*[@id="projectRegister-additionalOpt"]/div[6]/div[2]/div[2]/div/div[3]/ul/li[1]').click()

    else:
        driver.find_element(By.XPATH,'//*[@id="projectRegister-additionalOpt"]/div[6]/div[2]/div[2]/div/div[3]/ul/li[2]').click()

    driver.find_element(By.XPATH,"(.//*[normalize-space(text()) and normalize-space(.)='취약점 점검'])[1]/following::span[1]").click()
    time.sleep(0.5)

    if type.lower() == 'static':
        driver.find_element(By.XPATH,'//*[@id="projectRegister-additionalOpt"]/div[6]/div[5]/div[2]/div/div[3]/ul/li[1]').click()

    else:
        driver.find_element(By.XPATH,'//*[@id="projectRegister-additionalOpt"]/div[6]/div[5]/div[2]/div/div[3]/ul/li[2]').click()

    driver.find_element(By.ID, 'projectBtn').click()


def unzip_zipfile(case_id, type):
    filename = case_id + '_' + type + '.zip'

    directory_file = os.getcwd() + "\\" + 'report'
    excel_file_path = os.path.join(directory_file, filename)

    with zipfile.ZipFile(excel_file_path, 'r') as zip_ref:
        zip_ref.extractall(directory_file)


def reupload_flie(driver, case_id, type):
    # 프로젝트 결과 다운로드 - 다운로드 된 압축 파일 해제
    download_project(driver, case_id, type)
    unzip_zipfile(case_id, type)

    close_episode(driver, case_id)

    # 다운로드한 파일를 다시 결과 업로드
    script_directory = os.path.dirname(os.path.abspath(__file__))

    file_name = '*_' + type + '_RESULT.xml'

    zip_name = case_id + "_" + type + ".zip"
    file_path = os.path.join(script_directory, 'report', file_name)
    zip_path = os.path.join(script_directory, 'report', zip_name)

    matching_files = glob.glob(file_path)

    # 프로젝트 개요 페이지 이동
    move_menu.move_progress(driver, case_id)
    time.sleep(1)

    # 신뢰성 시험 결과 클릭
    driver.find_element(By.XPATH,'//div[2]/div[2]/div[2]/button').click()
    time.sleep(1)

    # 업로드 클릭
    driver.find_element(By.XPATH,"//a[contains(text(),'업로드')]").click()

    # 업로드할 도구 클릭
    driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='STATIC']/parent:: *").click()
    if type.lower() == 'static':
        driver.find_element(By.XPATH,'//div[14]/div/div/div/div[2]/div/div/div/div[3]/ul/li/span').click()

    elif type == 'Code Inspector' or type == 'CI':
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='Code Inspector']/parent::*").click()

    elif type == 'SNIPER' or type == 'SN':
        driver.find_element(By.XPATH,"//div[@id='projectDetailOverView-uploadDownloadModal']/div/div/div[2]/div/div/div/div[3]/ul/li[3]/span").click()

    elif type == 'Controller Tester' or type == 'CT':
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='CT']/parent::*").click()

    else:
        driver.find_element(By.XPATH,"//*/text()[normalize-space(.)='COVER EE/SE']/parent::*").click()

    # 파일 중 하나를 선택 (여기서는 첫 번째 파일 선택)
    if matching_files:
        file_to_upload = matching_files[0]

    driver.find_element(By.ID, "projectDetailOverView-resultFile").send_keys(file_to_upload)

    driver.find_element(By.ID, 'projectDetailOverView-uploadBtn').click()

    while True:
        try:
            # uploadResultList 엘리먼트 찾기
            upload_result_list = driver.find_element(By.ID, 'uploadResultList')

            # uploadResultList 엘리먼트에 text 값이 있으면 루프 종료
            if upload_result_list.text.strip():
                break

            time.sleep(1)

        # 못찾으면 false
        except NoSuchElementException:
            False

    time.sleep(5)

    # 신뢰성 시험 결과 창 닫기
    driver.find_element(By.XPATH,'//*[@id="projectDetailOverView-uploadDownloadModal"]/div/div/div[1]/button/img').click()

    # 프로젝트 상태가 진행중일 경우에 while문 종료.
    progress_status = driver.find_element(By.XPATH,'//*[@id="projectSideBar"]/div/span[1]')
    while True:
        if progress_status.text == "진행중":
            break

        time.sleep(1)

    # 프로젝트 결과 다운로드한 압축파일과 xml 파일 제거
    file_delete(file_to_upload)
    file_delete(zip_path)


# 문서간 일치 관계표 클릭 - 각각의 문서 식별자 확인 후 합산
def document_equal_relationships(driver):
    global identifier_sum
    currentWindow = driver.window_handles

    # 문서간 일치 관계표 클릭
    time.sleep(5)
    driver.find_element(By.ID, 'matchingTableClick').click() # Fail 자주 나는 구간
    time.sleep(3)

    # 현재 창을 문서간 일치 관계표 창으로 전환
    newWindow = driver.window_handles
    newWindow = list(set(newWindow) - set(currentWindow))[0]
    driver.switch_to.window(newWindow)

    # 각각의 문서 식별자 확인
    identifier1 = driver.find_elements(By.CSS_SELECTOR, 'div.left-no-match')
    identifier2 = driver.find_elements(By.CSS_SELECTOR, 'div.up-no-match')
    identifier_sum = len(identifier1) + len(identifier2)

def last_page(driver, result_type):
    if result_type == "coding":
        result_type_id = "codingViolationDetailByFile"
        next_page = "codingViolationDetailByFileTable_next"
        paginate_id = "codingViolationDetailByFileTable_paginate"
        length_id = "codingViolationDetailByFileTable_length"

    elif result_type == "weak" :
        result_type_id = "rteViolationDetailByFile"
        next_page = "rteViolationDetailByFileTable_next"
        paginate_id = "rteViolationDetailByFileTable_paginate"
        length_id = "rteViolationDetailByFileTable_length"

    elif result_type == "file" :
        result_type_id = "coverageByFile"
        next_page = "coverageByFileTable_next"
        paginate_id = "coverageByFileTable_paginate"
        length_id = "coverageByFileTable_length"

    else:
        result_type_id = "function"
        next_page = "coverageByFunctionTable_next"
        paginate_id = "coverageByFunctionTable_paginate"
        length_id = "coverageByFunctionTable_length"

    driver.find_element(By.ID, length_id).click()

    # '100'을 선택
    driver.find_element(By.XPATH,f'//*[@id="{length_id}"]/label/select/option[5]').click()

    time.sleep(10)
    common_css_selector = f'div#{paginate_id} span a:last-child'
    last_page_element = driver.find_element(By.CSS_SELECTOR, common_css_selector)
    last_page = int(last_page_element.text)

    date_result = []

    if last_page >= 1:
        for _ in range(last_page):
            for n in range(1, 101):
                # 데이터 수집
                time.sleep(0.2)
                xpath = f'//*[@id="{result_type_id}"]/tr[{n}]'

                try:
                    # XPath로 해당 tr 요소를 찾고 텍스트 값을 리스트에 추가
                    tr_text = driver.find_element(By.XPATH,xpath).text
                    date_result.append(tr_text)
                except:
                    # 더 이상 tr을 찾을 수 없으면 종료
                    break

            # 다음 페이지로 이동
            time.sleep(1)
            driver.find_element(By.ID, next_page).click()

    return date_result

# 기술 문서 검증 -> documents 폴더 안 파일 업로드
# subfolder_name는 일반 검증 위배 폴더 한에서 세부 폴더명 적어주시면 됩니다. ex)SRS_위배
def upload_document(driver,forder_name,file_name ,file_type,subfolder_name=None):
    #파일 경로 지정
    directory = os.getcwd()
    if subfolder_name is not None:
        file_path = os.path.join(directory, 'documents', forder_name, subfolder_name, file_name)
    else:
        file_path = os.path.join(directory, 'documents', forder_name, file_name)

    # 기술 문서 검증 - 파일 선택 클릭
    driver.find_element(By.ID, file_type.lower() + 'InputFile').send_keys(file_path)
    time.sleep(1.5)

    wait = WebDriverWait(driver, 5)
    #upload_button = wait.until(EC.element_to_be_clickable((By.ID,file_type.lower()+'UploadBtn')))
    #upload_button.click()
    time.sleep(2)

def download_document(driver,file_type):
    if file_type == 1 or file_type == 'SRS':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[1]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[1]/div/div/div/div[2]/div/a[contains(., "다운로드")]').click()
    elif file_type == 3 or file_type == 'SDD':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[3]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[3]/div/div/div/div[2]/div/a[contains(., "다운로드")]').click()
    elif file_type == 5 or file_type == 'STP':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[5]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[5]/div/div/div/div[2]/div/a[contains(., "다운로드")]').click()
    elif file_type == 7 or file_type == 'STD':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[7]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[7]/div/div/div/div[2]/div/a[contains(., "다운로드")]').click()
    elif file_type == 9 or file_type == 'STR':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[9]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[9]/div/div/div/div[2]/div/a[contains(., "다운로드")]').click()
    elif file_type == 11 or file_type == 'SIP':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[11]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[11]/div/div/div/div[2]/div/a[contains(., "다운로드")]').click()
    elif file_type == 13 or file_type == 'SPS':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[13]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[13]/div/div/div/div[2]/div/a[contains(., "다운로드")]').click()
    time.sleep(1)

def delete_document(driver,file_type):
    if file_type == 1 or file_type == 'SRS':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[1]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[1]/div/div/div/div[2]/div/a[3]').click()
    elif file_type == 3 or file_type == 'SDD':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[3]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[3]/div/div/div/div[2]/div/a[3]').click()
    elif file_type == 5 or file_type == 'STP':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[5]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[5]/div/div/div/div[2]/div/a[3]').click()
    elif file_type == 7 or file_type == 'STD':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[7]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[7]/div/div/div/div[2]/div/a[3]').click()
    elif file_type == 9 or file_type == 'STR':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[9]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[9]/div/div/div/div[2]/div/a[3]').click()
    elif file_type == 11 or file_type == 'SIP':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[11]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[11]/div/div/div/div[2]/div/a[3]').click()
    elif file_type == 13 or file_type == 'SPS':
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[13]/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="projectDocVerificationVue"]/div[2]/div/div[13]/div/div/div/div[2]/div/a[3]').click()


# 설정 -> 프로젝트 설정 -> 필수 정보 -> 개정안 선택 -> 저장
def select_amendment(driver,case_id,amendent_year):
    # 설정 -> 프로젝트 설정 -> 추가 정보
    move_menu.move_projectSetting(driver, case_id)
    time.sleep(3)

    # 개정안 선택
    driver.find_element(By.XPATH,'//*[@id="projectRegister-requiredOpt"]/div[2]/div/div[2]/div/div/div[2]').click()

    if amendent_year == '2018':
        driver.find_element(By.XPATH,'//span[contains(text(), "방사청 제2018-7호")]').click()

    elif amendent_year == '2020':
        driver.find_element(By.XPATH,'//span[contains(text(), "방사청 제2020-1호")]').click()

    elif amendent_year == '2022':
        driver.find_element(By.XPATH,'//span[contains(text(), "방사청 제2022-6호")]').click()

    elif amendent_year == '2024':
        driver.find_element(By.XPATH,'//span[contains(text(), "방사청 제2024-6호")]').click()

    driver.find_element(By.ID, 'projectBtn').click()
    time.sleep(2)

# 기술 문서 -> 기술 문서 검증 -> 검증하기 버튼 누른 후 완료 대기
def click_inspectionBtn_wait(driver,file_type):
    if file_type == 1 or file_type == 'SRS':
        file_type = 1
    elif file_type == 3 or file_type == 'SDD':
        file_type = 3
    elif file_type == 5 or file_type == 'STP':
        file_type = 5
    elif file_type == 7 or file_type == 'STD':
        file_type = 7
    elif file_type == 9 or file_type == 'STR':
        file_type = 9
    elif file_type == 11 or file_type == 'SIP':
        file_type = 11
    elif file_type == 13 or file_type == 'SPS':
        file_type = 13


# 검증하기 버튼 클릭
    driver.find_element(By.XPATH,f'//*[@id="projectDocVerificationVue"]/div[2]/div/div[{file_type}]/div/div/div/div[1]/button').click()

    # 검증 완료될때까지 대기
    #wait = WebDriverWait(driver, 60)
    #element = wait.until(EC.element_to_be_clickable((By.XPATH, "//*/text()[normalize-space(.)='검증하기']/parent::*")))
    #time.sleep(3)



# 기술 문서 -> 기술 문서 검증 결과 -> 일반 위배 검색
def search_diffrent_data(driver,search_data):
    driver.find_element(By.XPATH,"(//input[@type='search'])[3]").click()
    driver.find_element(By.XPATH,"(//input[@type='search'])[3]").clear()
    driver.find_element(By.XPATH,"(//input[@type='search'])[3]").send_keys(search_data)
    driver.find_element(By.XPATH,"(//input[@type='search'])[3]").send_keys(Keys.ENTER)
    time.sleep(1)




# 기술문서 -> 일반 검증 결과창 첫 항목에 search_text가 포함될때까지 상세 내용의 화살표 버튼 클릭
def click_content_text_arrow_btn(driver,search_text):
    while True:

        # 상세 내용 화살표 버튼 클릭
        arrow_button = driver.find_element(By.XPATH,'//*[@id="GeneralInspectionResultTable_wrapper"]/div[5]/div[1]/div/table/thead/tr/th[4]/div')
        arrow_button.click()

        # 일반 검증 결과창에 search_text가 발견될때까지 계속 클릭
        try:
            footer_text = WebDriverWait(driver, 10).until(
                EC.text_to_be_present_in_element(
                    (By.XPATH, '//*[@id="GeneralInspectionResultTable"]/tbody/tr[1]/td[4]'), search_text))
            break
        except Exception:
            continue
        time.sleep(1)

# 정적 시험 -> 파일별 결함 상세 결과 -> 검색 후 최상위 파일 예외 처리
def static_exception(driver, type, search_data, img = None):
    search_static_file_result(driver,type, search_data)
    if type == "코딩":
        id = "codingViolationDetailByFile"
    elif type == "취약점":
        id = "rteViolationDetailByFile"
    else:
        id ="rteViolationDetailByFile"
    # 정적 시험 -> 파일별 결함 상세 결과 -> 최상위 파일 선택
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/tr[1]').click()
    time.sleep(1)
    # 정적 시험 예외 사유 작성 -> 진단 메시지 사유 작성
    textarea = driver.find_element(By.XPATH,
        '//*[@id="ProjectDetailStaticExceptionModal"]/div/div/div[2]/div[1]/div/table[2]/tbody/tr[2]/td[2]/textarea')
    textarea.send_keys('정적 시험 예외 처리 test')

    message_xpath = '//*[@id="ProjectDetailStaticExceptionModal"]/div/div/div[2]/div[1]/div/table[2]/tbody/tr[1]/td[2]/textarea'
    message_element = driver.find_element(By.XPATH ,message_xpath)  # 웹 요소 찾기
    project_img = message_element.get_attribute('class')
    if 'is-invalid' in project_img:
        message_element.send_keys("진단 메시지")

    if img:
        driver.find_element(By.ID, 'projectDetailStaticException-resultFile').send_keys(img)

    # 진단 메시지 저장 후 닫기 버튼 클릭
    driver.find_element(By.XPATH,'//*[@id="ProjectDetailStaticExceptionModal"]//a[contains(text(), "저장")]').click()
    time.sleep(2)
    driver.find_element(By.XPATH,'//*[@id="ProjectDetailStaticExceptionModal"]//button[contains(text(), "닫기")]').click()
    time.sleep(2)

def unzip_report_zipfile(filename):
    filename = filename + '.zip'

    directory_file = os.getcwd() + "\\" + 'report'
    excel_file_path = os.path.join(directory_file, filename)

    with zipfile.ZipFile(excel_file_path, 'r') as zip_ref:
        zip_ref.extractall(directory_file)



def extract_table_values(doc, table_index):
    if table_index < 0 or table_index >= len(doc.tables):
        return "Invalid table index"

    table = doc.tables[table_index]
    table_data = []

    # 표 데이터 추출
    for row in table.rows:
        row_data = []
        for cell in row._element.xpath('.//w:tc'):
            # 병합된 셀도 모두 텍스트로 가져오기
            text = cell.xpath('.//w:t')
            row_data.append("".join([t.text for t in text if t.text]))
        table_data.append(row_data)

    return table_data

    return table_data

def add_user(driver, case_id):
    # 1. 로그인 페이지 -> 사용자 등록 클릭
    account_delete(driver, case_id)
    # 로그아웃 하기
    driver.find_element(By.ID, "TopNavBarUserInfoDropDownBtn").click()
    driver.find_element(By.XPATH,'//a[contains(text(), "로그아웃")]').click()
    time.sleep(1)

    driver.find_element(By.XPATH,'//a[contains(text(), "사용자 등록")]').click()

    # 2. 필수 항목(아이디 , 비밀번호 , 사용자 이름) 입력 후 [사용자 등록 신청] 버튼 클릭
    driver.find_element(By.XPATH,"//div[1]/div[1]/div[2]/input").send_keys(case_id)
    driver.find_element(By.XPATH,"//div[1]/div[2]/div[2]/input").send_keys(case_id)
    driver.find_element(By.XPATH,"//div[2]/div[1]/div[2]/input").send_keys("qwer1234!!")
    driver.find_element(By.XPATH,"//div[2]/div[2]/div[2]/input").send_keys("qwer1234!!")
    time.sleep(0.5)

    # 사용자 등록 신청 선택
    driver.find_element(By.ID, "userManageRegister-saveBtn").click()
    time.sleep(2)

# 산출물 카테고리(신뢰성, SW, KOLAS 선택)

def Report_category(driver, type):
    driver.find_element(By.XPATH, '//*[@id="reportGenerateVue"]/div[2]/div[2]/div[1]/div/div[1]').click()

    if type == '신뢰성':
        driver.find_element(By.XPATH,'//span[contains(text(), "신뢰성 문서")]').click()
    elif type == "SW":
        driver.find_element(By.XPATH,'//span[contains(text(), "SW 문서")]').click()
    elif type == "KOLAS":
        driver.find_element(By.XPATH,'//span[contains(text(), "KOLAS 문서")]').click()

def Report_type(driver, type): #프로젝트 -> 산출물 양식 선택
    driver.find_element(By.XPATH,'//*[@id="reportGenerateVue"]/div[2]/div[2]/div[2]/div/div[2]').click()

    if type == '엑셀' or type == 'xlsx':
        driver.find_element(By.XPATH,'//span[contains(text(), "엑셀")]').click()
    elif type == "한글 개별" or type == "hwp[개별]":
        driver.find_element(By.XPATH,'//span[contains(text(), "한글 개별")]').click()
    elif type == "한글 통합" or type == "hwp[통합]":
        driver.find_element(By.XPATH,'//span[contains(text(), "한글 통합")]').click()
    else:
        driver.find_element(By.XPATH,'//span[contains(text(), "워드")]').click()


def Report_type_SW(driver, type): #프로젝트 -> 산출물 양식 선택
    driver.find_element(By.XPATH,'//*[@id="reportGenerateVue"]/div[2]/div[2]/div[2]/div/div[2]').click()

    if type == '엑셀' or type == 'xlsx':
        driver.find_element(By.XPATH,'//*[@id="reportGenerateVue"]/div[2]/div[2]/div[2]/div/div[3]/ul/li[3]').click()
    elif type == "한글 개별" or type == "hwp[개별]":
        driver.find_element(By.XPATH,'//*[@id="reportGenerateVue"]/div[2]/div[2]/div[2]/div/div[3]/ul/li[1]').click()
    else:
        driver.find_element(By.XPATH,'//*[@id="reportGenerateVue"]/div[2]/div[2]/div[2]/div/div[3]/ul/li[2]').click()

# 특정 산출물을 선택(클릭)하는 모듈
def Report_select(driver, reportname):
    driver.find_element(By.XPATH, f"//*[@id='reportDataTable']/tbody//td[text()='{reportname}']").click()
    time.sleep(1)

# [산출물 생성] 버튼 클릭 후 산출물 생성이 완료 될 때 까지 대기
def Report_generate(driver):
    time.sleep(3)
    driver.find_element(By.XPATH, '//*[@id="btnaReportGenerate"]/div[contains(text(),"산출물 생성")]').click()
    wait = WebDriverWait(driver, 60)
    wait.until(EC.invisibility_of_element_located((By.XPATH, '//*[@id="generatedReportsTable"]/tbody/tr/td[contains(text(),"산출물을 생성중입니다.")]')))
    time.sleep(3)

# 생성 된 산출물 중 지적한 텍스트가 포함된 산출물을 클릭하여 다운로드
def Report_download(driver, report):
    driver.find_element(By.XPATH, f"//*[@id='generatedReportsTable']/tbody/tr/td[contains(text(),'{report}')]/div").click()
    time.sleep(5)


def search_static_file_result(driver,type, search_data):
    if type == "코딩":
        id = "codingViolationDetailByFileTable_filter"
    elif type == "취약점":
        id = "rteViolationDetailByFileTable_filter"
    else:
        id ="secureCodingViolationDetailByFileTable_filter"

    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label[2]/input').click()
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label[2]/input').clear()
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label[2]/input').send_keys(search_data)
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label[2]/input').send_keys(Keys.ENTER)
    time.sleep(2)


def search_dynamic_file_result(driver,type, search_data):
    if type == "파일":
        id = "coverageByFileTable_filter"
    else:
        id ="coverageByFunctionTable_filter"

    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').click()
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').clear()
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').send_keys(search_data)
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').send_keys(Keys.ENTER)
    time.sleep(1)


def search_metric_file_result(driver, search_data):
    id = "sourceMetricDetailTable_filter"
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').click()
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').clear()
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').send_keys(search_data)
    driver.find_element(By.XPATH,f'//*[@id="{id}"]/label/input').send_keys(Keys.ENTER)
    time.sleep(1)

#규칙설정에 위배무시 - 모두무시 후 정적검증 TAB(코딩/취약점/보안성)에서 위배가 정상적으로 체크 되는지 확인 한 모듈
def ruleset_Violation(driver, case_id, type, ruletpye, search_data, violation):
    move_menu.move_ruleset(driver, case_id)
    if type == '코딩':
        n = '1'
    elif type == '취약점':
        n = '2'
    elif type == '보안성':
        n = '3'

    driver.find_element(By.XPATH, '//*[@id="projectRuleSettingVue"]/div[2]/div[1]/div[1]/div[1]').click()
    driver.find_element(By.XPATH, f'//*[@id="projectRuleSettingVue"]/div[2]/div[1]/div[1]/div[3]/ul/li[{n}]').click()

    driver.find_element(By.XPATH, '//*[@id="projectRuleSettingVue"]/div[2]/div[1]/div[2]/div[1]').click()
    time.sleep(0.5)
    driver.find_element(By.XPATH, f"//span/span[text()='{ruletpye}']").click()
    time.sleep(2)

    driver.find_element(By.XPATH, '//*[@id="projectRuleDataTable_filter"]/label/input').send_keys(search_data)
    driver.find_element(By.XPATH, '//*[@id="projectRuleDataTable_filter"]/label/input').send_keys(Keys.ENTER)
    time.sleep(4)

    driver.find_element(By.XPATH, '//*[@id="projectRuleDataTable"]/tbody/tr/td[1]').click()
    time.sleep(3)

    driver.find_element(By.XPATH, f'//*[@id="projectRuleIgnoreModal"]/div/div/div[2]//div[2]/div/div[1]').click()
    if violation == '모두무시':
        n1 = '1'
    elif violation == '재사용함수':
        n1 = '2'
    elif violation == '해제':
        n1 = '3'
    driver.find_element(By.XPATH, f'//*[@id="projectRuleIgnoreModal"]/div/div/div[2]//div[2]/div/div[3]/ul/li[{n1}]').click()
    time.sleep(2)
    driver.find_element(By.ID, 'ignoreSaveBtn').click()

def add_KnowledgeHub_rule(driver, case_id, ruleset, type, rule=False):
    # type에 따라 올바른 메뉴 이동 함수와 XPath 값을 설정
    if type == "static":
        move_menu.move_knowledgehub_static_rule(driver)
        time.sleep(1)
        textbox_xpath = '//div[2]/div/div[6]/div[2]/div[2]/div'
        button_xpath = '//*[@id="knowledgeHubRuleInfoViewVue"]/div[2]/div/div[10]/div[2]/div/button'

    elif type == "metric":
        move_menu.move_knowledgehub_metric_rule(driver)
        time.sleep(1)
        textbox_xpath = '//div[4]/div[2]/div[2]/div'
        button_xpath = '//*[@id="knowledgeHubRuleInfoViewVue"]/div[2]/div/div[8]/div[2]/div/button'
    else:
        raise ValueError("Invalid type provided: must be 'static' or 'metric'")

    # 규칙 생성 및 제목, 룰셋 입력
    driver.find_element(By.XPATH, '//div[2]/button/span').click()
    driver.find_element(By.XPATH, '//div[2]/div[1]/div[2]/input').send_keys(case_id)
    driver.find_element(By.XPATH, '//div[2]/div[2]/div[2]/input').send_keys(ruleset)

    if type == 'static':
        # 규칙 입력
        driver.find_element(By.XPATH, '//div[3]/div[2]/input').send_keys(rule)

    # 텍스트 박스를 클릭하여 포커스를 설정하고, 텍스트를 입력
    textbox_element = driver.find_element(By.XPATH, textbox_xpath)
    textbox_element.click()

    # 현재 텍스트를 지우기 위해 Ctrl+A (전체 선택) 후, Backspace 키 사용
    textbox_element.send_keys(Keys.CONTROL + 'a')
    textbox_element.send_keys(Keys.BACKSPACE)
    textbox_element.send_keys(case_id)

    # [작성] 버튼 클릭
    driver.find_element(By.XPATH, button_xpath).click()


def delete_KnowledgeHub(driver, case_id, type):
    move_menu.search_KnowledgeHub(driver, case_id, type)

    if type == "static":
        n = '10'
        n1 = '13'
    elif type == "static_exercise":
        n = '8'
        n1 = '9'
    elif type == "metric":
        n = '8'
        n1 = '11'
    elif type == "metric_exercise":
        n = '8'
        n1 = '9'
    elif type == "setting":
        n = '8'
        n1 = '10'
    elif type == "trouble":
        n = '8'
        n1 = '10'
    elif type == "qna":
        n = '6'
        n1 = '8'
    else:
        raise ValueError(f"Unsupported type: {type}")

    # 삭제 버튼 클릭
    driver.find_element(By.XPATH, f'//div[{n1}]/div[2]/div/button').click()
    time.sleep(1)
    # 확인 버튼 클릭
    driver.find_element(By.XPATH, f'//div[{n}]/div[2]/div/button[2]').click()
    time.sleep(1)

def add_KnowledgeHub_exercise(driver, case_id, type, badtext, goodtext, in_project=False):
    if in_project:
        move_menu.search_KnowledgeHub(driver, case_id, type)
        driver.find_element(By.XPATH, '//*[@id="exampleTable"]/form/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="knowledgeTable_filter"]/label[1]/button').click()
    else:
        if type == 'static':
            move_menu.move_knowledgehub_static_exercise(driver)
        elif type == 'metric':
            move_menu.move_knowledgehub_metric_exercise(driver)

        driver.find_element(By.XPATH, '//*[@id="knowledgeHubTableboardVue"]/div[1]/div/div[2]/div[2]/button').click()
    time.sleep(1)
    driver.find_element(By.XPATH, '//div[1]/div[2]/input').send_keys(case_id)

    time.sleep(1)
    badcode = driver.find_element(By.XPATH, '//div[4]/div/div[1]/div[2]//textarea')
    goodcode = driver.find_element(By.XPATH, '//div[4]/div/div[2]/div[2]//textarea')

    driver.find_element(By.XPATH, '//div[1]/div[2]/div[6]/div[1]/div/div/div').click()
    badcode.send_keys(badtext)

    driver.find_element(By.XPATH, '//div[2]/div[2]/div[6]/div[1]/div/div/div').click()
    goodcode.send_keys(goodtext)

    textbox_element = '//div[2]/div/div[6]/div[2]/div[2]/div'
    driver.find_element(By.XPATH, textbox_element).send_keys(Keys.CONTROL + 'a')
    driver.find_element(By.XPATH, textbox_element).send_keys(Keys.BACKSPACE)
    driver.find_element(By.XPATH, textbox_element).send_keys(case_id)

    driver.find_element(By.XPATH, '//div[8]/div[2]/div/button').click()

def upload_document_knowledge(driver, forder_name,file_name ,subfolder_name=None):
    #파일 경로 지정
    directory = os.getcwd()
    if subfolder_name is not None:
        file_path = os.path.join(directory, 'documents', forder_name, subfolder_name, file_name)
    else:
        file_path = os.path.join(directory, 'documents', forder_name, file_name)

    # 기술 문서 검증 - 파일 선택 클릭
    driver.find_element(By.ID,'InputFile').send_keys(file_path)
    time.sleep(1.5)

    wait = WebDriverWait(driver, 5)
    time.sleep(2)

def download_know_file():
    #파일 경로 지정
    directory = os.getcwd()
    file_path = os.path.join(directory, 'report')
    matching_files = glob.glob(file_path)

    # 폴더 내의 모든 파일 이름 가져오기
    if os.path.exists(file_path):
        file_names = os.listdir(file_path)
        if file_names:
            first_file_name = file_names[0]
            full_file_path = os.path.join(file_path, first_file_name)  # 전체 경로 생성
            file_delete(full_file_path)  # 전체 경로를 전달

            return first_file_name
    return None  # 폴더가 없거나 파일이 없을 경우 None 반환

def add_KnowledgeHub_rule_file(driver, case_id, ruleset, type, forder_name,file_name,rule=False):
    # type에 따라 올바른 메뉴 이동 함수와 XPath 값을 설정
    if type == "static":
        move_menu.move_knowledgehub_static_rule(driver)
        textbox_xpath = '//div[2]/div/div[6]/div[2]/div[2]/div'
        button_xpath = '//div[10]/div[2]/div/button'

    elif type == "metric":
        move_menu.move_knowledgehub_metric_rule(driver)
        textbox_xpath = '//div[4]/div[2]/div[2]/div'
        button_xpath = '//div[8]/div[2]/div/button'
    else:
        raise ValueError("Invalid type provided: must be 'static' or 'metric'")

    # 규칙 생성 및 제목, 룰셋 입력
    driver.find_element(By.XPATH, '//div[2]/button/span').click()
    driver.find_element(By.XPATH, '//div[2]/div[1]/div[2]/input').send_keys(case_id)
    driver.find_element(By.XPATH, '//div[2]/div[2]/div[2]/input').send_keys(ruleset)

    if type == 'static':
        # 규칙 입력
        driver.find_element(By.XPATH, '//div[3]/div[2]/input').send_keys(rule)

    # 텍스트 박스를 클릭하여 포커스를 설정하고, 텍스트를 입력
    textbox_element = driver.find_element(By.XPATH, textbox_xpath)
    textbox_element.click()

    # 현재 텍스트를 지우기 위해 Ctrl+A (전체 선택) 후, Backspace 키 사용
    textbox_element.send_keys(Keys.CONTROL + 'a')
    textbox_element.send_keys(Keys.BACKSPACE)
    textbox_element.send_keys(case_id)

    upload_document_knowledge(driver, forder_name, file_name)
    driver.find_element(By.XPATH, '//label/div').click()

    driver.find_element(By.XPATH, button_xpath).click()


def add_KnowledgeHub_setting(driver, case_id, type):
    text1 = 'test'

    # Move to the correct menu based on the 'type' parameter
    if type == 'setting':
        move_menu.move_knowledgehub_tool_setting(driver)
    elif type == 'trouble':
        move_menu.move_knowledgehub_tool_trouble(driver)
    elif type == 'qna':
        move_menu.move_knowledgehub_qna(driver)
    else:
        raise ValueError("Invalid type provided. Use 'setting', 'trouble', or 'qna'.")

    # 공통 글 생성
    driver.find_element(By.XPATH, '//*[@id="knowledgeHubTableboardVue"]/div[1]/div/div[2]/div[2]/button').click()
    driver.find_element(By.XPATH, '//div[2]/div[1]/div[2]/input').send_keys(case_id)

    # 'setting' 및 'trouble' 타입에서만 'text1' 입력
    if type in ['setting', 'trouble']:
        driver.find_element(By.XPATH, '//div[3]/div[2]/input').send_keys(text1)

    # 'trouble' 타입에서만 추가 'text1' 입력
    if type == 'trouble':
        driver.find_element(By.XPATH, '//div[2]/div[2]/input').send_keys(text1)

    # Textbox 편집
    textbox_element = driver.find_element(By.XPATH, '//div[4]/div[2]/div[2]/div')
    textbox_element.click()
    textbox_element.send_keys(Keys.CONTROL + 'a')
    textbox_element.send_keys(Keys.BACKSPACE)
    textbox_element.send_keys(case_id)

    # 'qna' 타입의 경우, 다른 경로로 이동하여 작성
    if type == 'qna':
        driver.find_element(By.XPATH, '//div[6]/div[2]/div/button').click()
    else:
        # 'setting' 및 'trouble' 타입에서 작성
        driver.find_element(By.XPATH, '//div[8]/div[2]/div/button').click()


def add_KnowledgeHub_setting_file(driver, case_id, type, folder_name, file_name):
    # type에 따라 올바른 메뉴 이동 함수 호출
    text1 = 'test'

    if type == 'setting':
        move_menu.move_knowledgehub_tool_setting(driver)
    elif type == 'trouble':
        move_menu.move_knowledgehub_tool_trouble(driver)
    else:
        raise ValueError("Invalid type provided. Use 'setting' or 'trouble'.")

    # 공통 입력 필드 처리
    driver.find_element(By.XPATH, '//*[@id="knowledgeHubTableboardVue"]/div[1]/div/div[2]/div[2]/button').click()
    driver.find_element(By.XPATH, '//div[2]/div[1]/div[2]/input').send_keys(case_id)
    driver.find_element(By.XPATH, '//div[3]/div[2]/input').send_keys(text1)

    # 'trouble'의 경우에만 추가 입력 필드 처리
    if type == 'trouble':
        driver.find_element(By.XPATH, '//div[2]/div[2]/input').send_keys(text1)

    # 텍스트 박스 처리
    textbox_element = driver.find_element(By.XPATH, '//div[4]/div[2]/div[2]/div')
    textbox_element.click()
    textbox_element.send_keys(Keys.CONTROL + 'a')
    textbox_element.send_keys(Keys.BACKSPACE)
    textbox_element.send_keys(case_id)

    # 파일 업로드
    upload_document_knowledge(driver, folder_name, file_name)

    # 확인란 클릭 및 제출
    driver.find_element(By.XPATH, '//label/div').click()
    driver.find_element(By.XPATH, '//div[8]/div[2]/div/button').click()

def add_globalSetting_dir(driver, case_id, dirname):
    move_menu.move_globalSetting_scm(driver)
    # dir 존재시 삭제
    try:
        element = driver.find_element(By.XPATH, '//div[2]/div/div/form/div/input[1]')
        if element:
            print("요소가 존재합니다. 삭제 작업을 수행합니다.")
            # 삭제 아이콘 클릭
            driver.find_element(By.XPATH, "//div[2]/div/div/form/div/img").click()
            time.sleep(1)
            try:
                # 특정 요소 찾기
                toast_element = driver.find_element(By.ID, "toastBottomCenter_COMMON_SUCCESS0")

                # class 속성 값 가져오기
                class_attribute = toast_element.get_attribute('class')

                # 'show' 클래스가 포함되어 있는지 확인
                if 'show' in class_attribute:
                    projectname = driver.find_element(By.XPATH, '(//div[@class="toast-msg"])[1]').text
                    lines = projectname.split('\n')
                    project1 = lines[2]
                    delete_project(driver, project1)
                    move_menu.move_globalSetting_scm(driver)
                    driver.find_element(By.XPATH, "//div[2]/div/div/form/div/img").click()
            except:
                print("프로젝트 삭제 안함")

            driver.find_element(By.ID, 'btn-scm').click()
    except:
        print("요소를 찾을 수 없습니다.")

    # 2. [+항목추가] > 별칭과 경로 입력 후 저장
    driver.find_element(By.XPATH, '//*[@id="scmSetting"]/div[2]/div/a/font').click()
    time.sleep(1)
    driver.find_element(By.XPATH, '//div[2]/div/div/form/div/input[1]').send_keys(case_id)
    driver.find_element(By.XPATH, '//div[2]/div/div/form/div/input[2]').send_keys(dirname)
    driver.find_element(By.ID, 'btn-scm').click()

def user_login(driver, case_id):
    # VPES 서버 열기 및 로그인
    driver.get(addressLogin)
    driver.implicitly_wait(30)

    driver.find_element(By.ID, "username").send_keys(case_id)
    driver.find_element(By.ID, "pwd").send_keys("qwer1234!!")
    driver.find_element(By.XPATH, '//button[contains(text(), "로그인")]').click()
    time.sleep(2)
    log_setting.logger.info("The account is available.")

    try:
        login_check = driver.find_element(By.ID, "TopNavBarUserInfoDropDownBtn")
        result = login_check.is_displayed()
    except NoSuchElementException:
        result = False

def delete_CheckAssignmentInformation(driver, case_id, type = None):
    move_menu.move_CheckAssignmentInformation(driver)

    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="projectHistoryList"]/thead/tr/th[1]/a').click()
    driver.find_element(By.XPATH, '//*[@id="projectHistoryList"]/thead/tr/th[1]/div/div[3]/font').click()

    if type == None:
        print('전체 삭제')
    else:
        driver.find_element(By.XPATH, '//*[@id="projectHistoryList_filter"]/label/input').click()
        driver.find_element(By.XPATH, '//*[@id="projectHistoryList_filter"]/label/input').send_keys(case_id)
        driver.find_element(By.XPATH, '//*[@id="projectHistoryList_filter"]/label/input').send_keys(Keys.ENTER)

    time.sleep(3)
    common_css_selector = 'div#projectHistoryList_paginate span a:last-child'
    last_page_element = driver.find_element(By.CSS_SELECTOR, common_css_selector)
    last_page = int(last_page_element.text)

    if last_page >= 1:
        for _ in range(last_page):
            for n in range(1, 11):
                # 데이터 수집
                time.sleep(0.2)
                xpath = f'//*[@id="projectHistoryList"]/tbody/tr[{n}]'

                try:
                    # XPath로 해당 tr 요소를 찾고 텍스트 값을 리스트에 추가
                    element = WebDriverWait(driver, 3).until(
                        EC.presence_of_element_located((By.XPATH, xpath))
                    )
                    element.click()
                except:
                    # 더 이상 tr을 찾을 수 없으면 종료
                    break

            time.sleep(2)
            driver.find_element(By.XPATH, '//*[@id="projectHistoryList_wrapper"]/div[1]/button[3]').click()
            time.sleep(2)
            driver.find_element(By.ID, 'deleteY').click()


def click_and_search(driver, section_xpath, input_xpath, value):
    driver.find_element(By.XPATH, section_xpath).click()  # 섹션 클릭
    driver.find_element(By.XPATH, input_xpath).send_keys(value)  # 값 입력
    driver.find_element(By.XPATH, input_xpath).send_keys(Keys.ENTER)  # 엔터로 검색
    time.sleep(0.5)  # 필터 적용 후 대기 시간
    time.sleep(0.5)  # 필터 적용 후 대기 시간


def dynamic_exception(driver, type, search_data, img = None):
    search_dynamic_file_result(driver,type, search_data)
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="coverageByFunction"]/tr[1]').click()  # 파일별 결함 상세 테이블 임의 클릭.
    time.sleep(0.5)

    driver.find_element(By.XPATH,'//*[@id="ProjectDetailDynamicExceptionModal"]/div/div/div[2]/div[3]/div/table/tbody/tr/td[2]/textarea').send_keys("동적 시험 예외 사유 테스트")
    time.sleep(1)
    if img:
        driver.find_element(By.ID, 'projectDetailDynamicException-resultFile').send_keys(img)
    driver.find_element(By.XPATH,'//*[@id="ProjectDetailDynamicExceptionModal"]/div/div/div[4]/a[2]').click()  # 저장 버튼 클릭
    time.sleep(1)
    driver.find_element(By.XPATH,'//*[@id="ProjectDetailDynamicExceptionModal"]/div/div/div[4]/button').click()  # 닫기 버튼 클릭

#소스 코드 메트릭 점검 기준 변경
def metric_check_criteria(driver, case_id, type, num):
    move_menu.move_metric(driver, case_id)
    time.sleep(2)
    driver.find_element(By.ID, 'MetricSetting').click()

    if type == 'CC':
        n = '1'

    elif type == 'NCL':
        n = '2'

    elif type == 'NFP':
        n = '3'

    elif type == 'NCF':
        n = '4'

    elif type == 'NDF':
        n = '5'

    elif type == 'NECL':
        n = '6'

    input_field= driver.find_element(By.XPATH, f'//*[@id="ProjectDetailMetricStandardConfigModal"]/div/div/div[2]/div[{n}]/div[2]/input')
    input_field.send_keys(Keys.CONTROL + 'a')
    input_field.send_keys(Keys.BACKSPACE)
    input_field.send_keys(num)

    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailMetricStandardConfigModal"]/div/div/div[3]/button[2]').click()
    time.sleep(2)

# 소스 코드 메트릭 예외
def metric_exception(driver, search_data, type, img = None):
    search_metric_file_result(driver, search_data)

    if type == 'CC':
        n = '5'

    elif type == 'NCL':
        n = '6'

    elif type == 'NFP':
        n = '7'

    elif type == 'NCF':
        n = '8'

    elif type == 'NDF':
        n = '9'

    elif type == 'NECL':
        n = '10'

    # CC 내림차순 정렬
    driver.find_element(By.XPATH, f'//*[@id="sourceMetricDetailTable"]/thead/tr/th[{n}]').click()
    driver.find_element(By.XPATH, f'//*[@id="sourceMetricDetailTable"]/thead/tr/th[{n}]').click()
    time.sleep(1)
    # CC 가장 큰 값 클릭 하여 예외 작성
    driver.find_element(By.XPATH, f'//*[@id="sourceMetricDetail"]/tr[1]/td[{n}]').click()
    driver.find_element(By.XPATH, '//form/div/input[1]').send_keys("1")
    n = 2
    # 시작 줄 값 넣는 부분
    while True:
        try:
            # 텍스트가 존재하는지 확인할 요소
            element = driver.find_element(By.XPATH, '//*[@id="ProjectDetailMetricExceptionModal"]/div/div/div[2]/div[1]/div/div[4]/form/div/div')

            # 텍스트가 비어있지 않은 동안 반복
            if element.text.strip():
                input_field = driver.find_element(By.XPATH, '//form/div/input[1]')

                # 기존 값 지우기 (BACKSPACE)
                input_field.send_keys(Keys.CONTROL + 'a')
                input_field.send_keys(Keys.BACKSPACE)

                # 'n' 값을 입력
                input_field.send_keys(str(n))

                # 'n'을 증가
                n += 1

            else:
                print("텍스트가 비어있습니다. 반복 중지.")
                break  # 텍스트가 비어있으면 반복 중단

        except Exception as e:
            print(f"오류 발생: {e}")
            break

    # 끝 위치 값 넣는 부분
    driver.find_element(By.XPATH, '//form/div/input[2]').send_keys("1")
    input_field1 = driver.find_element(By.XPATH, '//form/div/input[2]')

    # 기존 값 지우기 (BACKSPACE)
    input_field1.send_keys(Keys.CONTROL + 'a')
    input_field1.send_keys(Keys.BACKSPACE)

    # 'n' 값을 입력
    input_field1.send_keys(str(n))
    # 대기 시간 (필요에 따라 조정)
    time.sleep(2)


    driver.find_element(By.XPATH, '//tbody/tr/td[2]/textarea').send_keys('메트릭 예외')

    if img:
        driver.find_element(By.ID, 'projectDetailMetricException-resultFile').send_keys(img)

    driver.find_element(By.XPATH, '//*[@id="ProjectDetailMetricExceptionModal"]/div/div/div[4]/a').click()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailMetricExceptionModal"]/div/div/div[4]/button').click()

# 소스 코드 메트릭 예외 - 사유를 사용자가 지정한 값으로 넣는 함수
def metric_exception_usertext(driver, search_data, type, text, img=None):
    search_metric_file_result(driver, search_data)

    if type == 'CC':
        n = '5'

    elif type == 'NCL':
        n = '6'

    elif type == 'NFP':
        n = '7'

    elif type == 'NCF':
        n = '8'

    elif type == 'NDF':
        n = '9'

    elif type == 'NECL':
        n = '10'

    # CC 내림차순 정렬
    driver.find_element(By.XPATH, f'//*[@id="sourceMetricDetailTable"]/thead/tr/th[{n}]').click()
    driver.find_element(By.XPATH, f'//*[@id="sourceMetricDetailTable"]/thead/tr/th[{n}]').click()
    time.sleep(1)
    # CC 가장 큰 값 클릭 하여 예외 작성
    driver.find_element(By.XPATH, f'//*[@id="sourceMetricDetail"]/tr[1]/td[{n}]').click()
    driver.find_element(By.XPATH, '//form/div/input[1]').send_keys("1")
    n = 2
    # 시작 줄 값 넣는 부분
    while True:
        try:
            # 텍스트가 존재하는지 확인할 요소
            element = driver.find_element(By.XPATH, '//*[@id="ProjectDetailMetricExceptionModal"]/div/div/div[2]/div[1]/div/div[4]/form/div/div')

            # 텍스트가 비어있지 않은 동안 반복
            if element.text.strip():
                input_field = driver.find_element(By.XPATH, '//form/div/input[1]')
                # 기존 값 지우기 (BACKSPACE)
                input_field.send_keys(Keys.CONTROL + 'a')
                input_field.send_keys(Keys.BACKSPACE)

                # 'n' 값을 입력
                input_field.send_keys(str(n))
                # 'n'을 증가
                n += 1

            else:
                print("텍스트가 비어있습니다. 반복 중지.")
                break  # 텍스트가 비어있으면 반복 중단
        except Exception as e:
            print(f"오류 발생: {e}")
            break

    # 끝 위치 값 넣는 부분
    driver.find_element(By.XPATH, '//form/div/input[2]').send_keys("1")
    input_field1 = driver.find_element(By.XPATH, '//form/div/input[2]')

    # 기존 값 지우기 (BACKSPACE)
    input_field1.send_keys(Keys.CONTROL + 'a')
    input_field1.send_keys(Keys.BACKSPACE)

    # 'n' 값을 입력
    input_field1.send_keys(str(n))
    # 대기 시간 (필요에 따라 조정)
    time.sleep(1)

    driver.find_element(By.XPATH, '//tbody/tr/td[2]/textarea').send_keys(text)

    if img:
        driver.find_element(By.ID, 'projectDetailMetricException-resultFile').send_keys(img)

    driver.find_element(By.XPATH, '//*[@id="ProjectDetailMetricExceptionModal"]/div/div/div[4]/a').click()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailMetricExceptionModal"]/div/div/div[4]/button').click()

def multirepository_setting(driver, case_id, num, type = None):
    move_menu.move_project_scm_setting(driver, case_id)
    url = ["https://github.com/YunNaRi/VPES_Autumation.git", "https://10.10.111.235/svn/VPES_Auto/study_sample", 'miro_vs2015']
    input_text = "test"  # 모든 입력값은 "Test"로 통일
    if type:
        if type =="GIT":
            x = 1
            num = 2
            time_num = 2
            i =0
        elif type == "SVN":
            x = 2
            num = 3
            time_num = 2
            i = 1
        else:
            x = 3
            num = 4
            time_num = 2
            i = 2
    else:
        x = 1
        i = 0
        time_num = num
    for n in range(x, num):
        driver.find_element(By.XPATH, '//*[@id="configurationManagementModal"]/div/div/div[2]/div[1]/div/div/div/div/div[2]').click()
        driver.find_element(By.XPATH, f'//div[2]/div[1]/div/div/div/div/div[3]/ul/li[{n}]').click()

    btn = '//*[@id="configurationManagementModal"]/div/div/div[3]/button[2]' #프로젝트 내 SCM 추가 저장 버튼
    scm_num = 2+time_num
    for n in range(3, scm_num):
        # scm 설정
        driver.find_element(By.XPATH,f'//div[{n}]/div[2]/div[2]/div[2]/div/input').send_keys(url[i]) #git url 입력
        # 인증 버튼 선택
        name = driver.find_element(By.XPATH,f'//div[{n}]/div[2]/div[2]/div[1]/div/div/div[2]/span').text
        if name == 'DIRECTORY':
            b = 3

        else:
            b = 5
            driver.find_element(By.XPATH, f"//div[{n}]/div[2]/div[2]/div[3]/div/input").send_keys(input_text) # ID 입력
            driver.find_element(By.XPATH, f"//div[{n}]/div[2]/div[2]/div[4]/div/input").send_keys(input_text) # PASSWORD 입력

        driver.find_element(By.XPATH, f"//div[{n}]/div[2]/div[2]/div[{b}]/div/button").click()
        time.sleep(1)

        # SCM 인증 결과 확인
        scm_img = driver.find_element(By.XPATH, f'//div[{n}]/div[2]/div[2]/div[2]/div/input').get_attribute('class')
        driver.find_element(By.XPATH,'//*[@id="projectRegister-scmUrl"]')

        if 'is-valid' not in scm_img:
            return False
        time.sleep(3)
        i = i + 1

    # 상단 저장 버튼 클릭
    driver.find_element(By.XPATH, btn).click()
    driver.find_element(By.XPATH, '//*[@id="configurationManagementModal"]/div/div/div[3]/button[1]').click()

# 프로젝트 - 그룹 정보 - 그룹 관리 창 진입 모듈
def click_group_management(driver, case_id):
    move_menu.move_groupinfo(driver, case_id)
    time.sleep(1)
    driver.find_element(By.XPATH, '//button/div[contains(text(),"그룹 관리")]').click()
    time.sleep(1)


def add_group(driver, case_id, type, num): #프로젝트에서 그룹 추가(CSC,CUS 추가)
    click_group_management(driver, case_id)
    num = int(num)

    if type == 'CSC':
        n = '1'
        x ='2'

    else:
        n = '2'
        x ='2'
        driver.find_element(By.XPATH, '//*[@id="groupManageEditModal"]/div/div/div[2]/div/div[1]/div[2]/ul/li[2]').click()
        time.sleep(0.5)
    for o in range(num):
        if o >= 1:
            driver.find_element(By.XPATH, f'//*[@id="groupManageEditModal"]/div/div/div[2]/div/div[{n}]/div[1]/button').click()
            driver.find_element(By.XPATH, f'//*[@id="groupManageEditModal"]/div/div/div[2]/div/div[{n}]/div[2]/ul/li[{x}]/div/input').send_keys(case_id+str(o))
            driver.find_element(By.XPATH, f'//*[@id="groupManageEditModal"]/div/div/div[2]/div/div[{n}]/div[2]/ul/li[{x}]/div/input').send_keys(Keys.ENTER)
        else:
            driver.find_element(By.XPATH, f'//*[@id="groupManageEditModal"]/div/div/div[2]/div/div[{n}]/div[1]/button').click()
            driver.find_element(By.XPATH, f'//*[@id="groupManageEditModal"]/div/div/div[2]/div/div[{n}]/div[2]/ul/li[{x}]/div/input').send_keys(case_id)
            driver.find_element(By.XPATH, f'//*[@id="groupManageEditModal"]/div/div/div[2]/div/div[{n}]/div[2]/ul/li[{x}]/div/input').send_keys(Keys.ENTER)
        x = int(x) + 1
        x = str(x)

    driver.find_element(By.XPATH, '//*[@id="groupManageEditModal"]//button/div[contains(text(),"저장")]').click()


def replace_text_in_word_at_position(file_name, target_text, replacement_text, occurrence, folder_name, tpye):
    # 문서 열기
    directory = os.getcwd()

    file_path = os.path.join(directory, 'documents', folder_name, file_name)

    doc = Document(file_path)
    count = 0

    if tpye == '글':
        # 문서 내의 모든 단락에서 텍스트 검색
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                count += 1
                # 지정된 위치의 텍스트만 치환 (3번째 값)
                if count == occurrence:
                    paragraph.text = paragraph.text.replace(target_text, replacement_text)
                    break

    elif tpye =='표':
        # 문서 내의 모든 표에서 텍스트 검색
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target_text in cell.text:
                        count += 1
                        # 지정된 위치의 텍스트만 치환
                        if count == occurrence:
                            cell.text = cell.text.replace(target_text, replacement_text)
                            break

    # 변경 사항 저장
    updated_file_name = 'new_' + file_name
    updated_file_path = os.path.join(directory, 'documents', folder_name, updated_file_name)
    doc.save(updated_file_path)
    print('변경 완료')

def interDocumentMatchMatrix(driver, doclist):
    driver.find_element(By.ID, 'matchingTableClick').click()
    # 새 창으로 이동
    new_window = driver.window_handles[1]  # 새로운 창의 핸들 가져오기
    driver.switch_to.window(new_window)     # 새 창으로 이동

    tooltip = []  # 툴팁 값을 저장할 리스트
    do_text = doclist  # 찾으려는 텍스트 리스트
    n = 0  # do_text의 인덱스

    while n < len(do_text):
        rows = driver.find_elements(By.TAG_NAME, "tr")

        for row_index in range(1, len(rows) + 1):
            # 첫 번째 셀의 텍스트를 가져옴
            text = driver.find_element(By.XPATH, f'//*[@id="matchingResultTable"]/tbody/tr[{row_index}]/td[1]').text

            if text == do_text[n]:  # 찾는 텍스트와 일치하면
                # 해당 행에서 특정 클래스 이름을 가진 div 요소들만 찾아 title 속성을 가져옴
                divs = driver.find_elements(By.XPATH, f'//*[@id="matchingResultTable"]/tbody/tr[{row_index}]//div[contains(@class, "left-no-match") or contains(@class, "up-no-match")]')

                for div in divs:
                    title = div.get_attribute("title")
                    if title:
                        tooltip.append(text + ', '+title)

                # do_text의 다음 인덱스로 이동하고, 외부 while 루프의 처음부터 다시 검색
                n += 1
                break  # 내부 for 루프를 빠져나가고 다시 처음부터 검색

    return tooltip


def LLM_settting_project(driver, case_id, type):

    move_menu.move_project(driver, case_id)
    driver.find_element(By.XPATH, '//span[contains(text(), "LLM 설정")]').click()

    if type.lower() == 'gemini':
        API_key = 'AIzaSyC02hwxjvuQydNJWIQPf1XLvBwo8SqekOo'
        type_xpath = driver.find_element(By.XPATH, '//span[contains(text(), "GEMINI")]')

    elif type.lower() == 'GPT':
        API_key = os.getenv("OPENAI_API_KEY")
        type_xpath = driver.find_element(By.XPATH, '//span[contains(text(), "GPT")]')

    # LLM 설정 사용
    driver.find_element(By.ID, 'aiUse').click()

    # LLM 타입 설정
    driver.find_element(By.XPATH, '//*[@id="TabAIRequest"]/div[2]/div[2]/div/div[2]').click()
    type_xpath.click()

    driver.find_element(By.XPATH, '//*[@id="TabAIRequest"]/div[3]/div[2]/textarea').send_keys(API_key)

    driver.find_element(By.XPATH, '//*[@id="aiSettingModal"]/div/div/div[4]/button[2]').click()
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="aiSettingModal"]/div/div/div[4]/button[1]').click()


def copy_project(driver, case_id, copy_id):
    move_menu.move_copy_project(driver)
    projectKey_check = driver.find_element(By.ID, "projectRegister-projectKey")
    projectKey_check.send_keys(case_id)
    time.sleep(1)
    driver.find_element(By.ID, 'projectRegister-duplicationBtn').click()
    time.sleep(1)
    # 프로젝트 명 중복 확인 결과 확인
    project_img = projectKey_check.get_attribute('class')
    if 'is-valid' not in project_img:
        driver.find_element(By.XPATH, '//*[@id="projectCopyModalVue"]/div/div/div[3]/div/button[1]').click()
        delete_project(driver, case_id)  # 중복 시 프로젝트 삭제
        copy_project(driver, case_id, copy_id)
        return True

    driver.find_element(By.XPATH, '//*[@id="projectCopyModalVue"]/div/div/div[2]/form[2]/div[2]/div/div[1]').click()
    driver.find_element(By.XPATH,f'//span[contains(text(), "{copy_id}")]').click()
    driver.find_element(By.ID, 'copyProjectBtn').click()

# 프로젝트 설정 - 보안성 시험(STATIC) 활성화 및 기간 설정
def activation_securecoding(driver,case_id):
    # 설정 -> 프로젝트 설정 -> 추가 정보
    move_menu.move_projectSetting(driver, case_id)
    driver.find_element(By.ID, 'projectRegister-additionalOpt-tab').click()
    time.sleep(1)

    # 보안성 시험 체크박스 클릭
    driver.find_element(By.ID,'projectRegister-secureCodingTestYN').click()

    # 보안성 시험 기간 설정
    driver.find_element(By.XPATH, '//*[@id="projectRegister-additionalOpt"]/div[7]/div[3]/div[2]/div/input').send_keys(
        '2020-01-01')
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="projectRegister-additionalOpt"]/div[7]/div[3]/div[3]/div/input').send_keys(
        '2030-01-01')
    time.sleep(1)

    # 저장
    driver.find_element(By.ID, 'projectBtn').click()
    time.sleep(2)


# [SW 구현]과 [SW 통합 및 시험]을 제외한 모든 프로세스 버튼의 XPATH 리스트 반환
def process_xpaths():
    return [
        "//*[@id='ProjectReliabilityProcessInfoVue']/div[2]/div[2]/button[1]",
        "//*[@id='ProjectReliabilityProcessInfoVue']/div[2]/div[2]/button[2]",
        "//*[@id='ProjectReliabilityProcessInfoVue']/div[2]/div[2]/button[3]",
        "//*[@id='ProjectReliabilityProcessInfoVue']/div[2]/div[2]/button[4]",
        "//*[@id='ProjectReliabilityProcessInfoVue']/div[2]/div[2]/button[5]",
        "//*[@id='ProjectReliabilityProcessInfoVue']/div[2]/div[2]/button[8]"
    ]

# 각 프로세스의 탭 클릭 -> [진행 체크 리스트] 탭 클릭
def click_processChecklist(driver, process_xpaths):
    processTap = driver.find_element(By.XPATH, process_xpaths)
    processTap.click() # 각 프로세스 탭 클릭
    time.sleep(1)

    driver.find_element(By.XPATH, "//div[4]/div/div/div/ul/li[2]/a").click()  # [진행 체크 리스트] 탭 클릭
    time.sleep(1)

# 각 프로세스의 진행 체크 리스트 마우스 오버
def mouseOver_checklist(driver):
    edit_icon = driver.find_element(By.XPATH,"//*[@id='checklist_contents']/div[1]/div/div/div[1]/div[1]/div[2]/div[2]")
    actions = ActionChains(driver)
    actions.move_to_element(edit_icon).perform()  # 마우스 오버
    time.sleep(1)

# [연필] 아이콘 클릭
def edit_checklist(driver):
    driver.find_element(By.XPATH,"//*[@id='checklist_contents']/div[1]/div/div/div[1]/div[1]/div[2]/div[3]/div[1]").click()  # [연필] 아이콘 클릭
    time.sleep(1)

    driver.find_element(By.XPATH, "//*[@id='btnGroupAddon']/div[1]").click()  # [체크] 버튼 클릭
    time.sleep(1)

# [휴지통] 아이콘 클릭
def delete_checklist(driver):
    edit_icon = driver.find_element(By.XPATH,"//*[@id='checklist_contents']/div[1]/div/div/div[1]/div[1]/div[2]/div[2]")
    actions = ActionChains(driver)
    actions.move_to_element(edit_icon).perform()  # 마우스 오버
    time.sleep(1)

    driver.find_element(By.XPATH,"//*[@id='checklist_contents']/div[1]/div/div/div[1]/div[1]/div[2]/div[3]/div[2]").click()
    time.sleep(1)

# 체크 리스트 체크 -> [진행 항목 확인 완료] 클릭
def check_exceptLastCheckbox(driver):
    checkboxes = driver.find_elements(By.XPATH,"//*[@id='checklist_contents']//input[@type='checkbox' and not (@id='srrCheck')]")
    for checkbox in checkboxes[:-1]:  # 마지막 체크 박스를 제외한 모든 체크 박스 클릭
        if not checkbox.is_selected():
            checkbox.click()
            time.sleep(0.2)

    driver.find_element(By.XPATH, "//*[@id='checklist_contents']/div[2]/div").click()  # [진행 항목 확인 완료] 버튼 클릭
    time.sleep(1)

# 전체 체크 -> [진행 항목 확인 완료] 클릭
def check_allCheckboxes(driver):
    checkboxes = driver.find_elements(By.XPATH, "//input[@type='checkbox' and not (@id='srrCheck')]")
    for checkbox in checkboxes:  # 모든 체크 박스 클릭
        if not checkbox.is_selected():
            checkbox.click()
            time.sleep(0.5)

    driver.find_element(By.XPATH, "//*[@id='checklist_contents']/div[2]/div").click()  # [진행 항목 확인 완료] 버튼 클릭
    time.sleep(1)

# 각 프로세스의 탭 클릭 -> [완료 산출물 관리] 탭 클릭
def click_outputManagement(driver, process_xpaths):
    processTap = driver.find_element(By.XPATH, process_xpaths)
    processTap.click() # 각 프로세스 탭 클릭
    time.sleep(1)

    driver.find_element(By.XPATH, "//*[@id='processContents']/div/ul/li[3]/a").click()  # [완료 산출물 관리] 탭 클릭
    time.sleep(1)

# 각 프로세스의 탭 클릭 -> [회의 이력] 탭 클릭
def click_meetingNotes(driver, process_xpaths):
    processTap = driver.find_element(By.XPATH, process_xpaths)
    processTap.click()  # 각 프로세스 탭 클릭
    time.sleep(1)

    driver.find_element(By.XPATH, "//*[@id='processContents']/div/ul/li[4]/a").click()  # [회의 이력] 탭 클릭
    time.sleep(1)

# Table의 헤더 컬럼을 클릭하여 해당 컬럼 기준으로 정렬
def sort_table(driver, table_id, column_number, order):
    if order == "ascending":
        if driver.find_element(By.XPATH, f'//*[@id="{table_id}"]//th[{column_number}]').get_attribute("aria-sort") == "ascending" :
            time.sleep(2)
        else:
            driver.find_element(By.XPATH, f'//*[@id="{table_id}"]//th[{column_number}]').click()
            time.sleep(2)
    elif order == "descending":
        if driver.find_element(By.XPATH, f'//*[@id="{table_id}"]//th[{column_number}]').get_attribute("aria-sort") == "descending" :
            time.sleep(2)
        elif driver.find_element(By.XPATH, f'//*[@id="{table_id}"]//th[{column_number}]').get_attribute("aria-sort") == "ascending" :
            driver.find_element(By.XPATH, f'//*[@id="{table_id}"]//th[{column_number}]').click()
            time.sleep(2)
        else:
            driver.find_element(By.XPATH, f'//*[@id="{table_id}"]//th[{column_number}]').click()
            driver.find_element(By.XPATH, f'//*[@id="{table_id}"]//th[{column_number}]').click()
            time.sleep(2)

# 파일 정보 페이지 - 파일 전체 선택 후 선택한 시험 항목만 비대상으로 변경
def filedata_change_nonobject(driver, examtype, text):
    driver.find_element(By.ID, 'fileInfoAllCheck').click()
    driver.find_element(By.XPATH, '//*[@id="projectInfoTable_wrapper"]/div[1]/button[2]').click()
    driver.find_element(By.XPATH, '//*[@id="projectInfoTable_wrapper"]/div[1]/div[2]/div/button[2]').click()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]/div/div/div[2]/form[2]/div/div/div[2]').click()
    driver.find_element(By.XPATH, '//span[contains(text(),"비대상")]').click()
    time.sleep(1)

    driver.find_element(By.ID, 'ciPartialTestTargetCheckBox').click()
    driver.find_element(By.ID, 'snPartialTestTargetCheckBox').click()
    driver.find_element(By.ID, 'scPartialTestTargetCheckBox').click()
    driver.find_element(By.ID, 'dynamicPartialTestTargetCheckBox').click()

    if examtype == "코딩규칙":
        driver.find_element(By.ID, 'ciPartialTestTargetCheckBox').click()
    elif examtype == "취약점":
        driver.find_element(By.ID, 'snPartialTestTargetCheckBox').click()
    elif examtype == "STATIC":
        driver.find_element(By.ID, 'ciPartialTestTargetCheckBox').click()
        driver.find_element(By.ID, 'snPartialTestTargetCheckBox').click()
    elif examtype == "보안성":
        driver.find_element(By.ID, 'scPartialTestTargetCheckBox').click()
    elif examtype == "동적":
        driver.find_element(By.ID, 'dynamicPartialTestTargetCheckBox').click()
    elif examtype == "all":
        driver.find_element(By.ID, 'ciPartialTestTargetCheckBox').click()
        driver.find_element(By.ID, 'snPartialTestTargetCheckBox').click()
        driver.find_element(By.ID, 'scPartialTestTargetCheckBox').click()
        driver.find_element(By.ID, 'dynamicPartialTestTargetCheckBox').click()

    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]//textarea').clear()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]//textarea').send_keys(text)
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]/div/div/div[3]/button[2]').click()

# 파일 정보 페이지 - 파일 전체 선택 후 검증 대상으로 변경
def filedata_change_object(driver, table_id, column_number, order):
    driver.find_element(By.ID, 'fileInfoAllCheck').click()
    driver.find_element(By.XPATH, '//*[@id="projectInfoTable_wrapper"]/div[1]/button[1]').click()
    driver.find_element(By.XPATH, '//button[contains(text(),"저장")]').click()

# 함수 정보 페이지 - 함수 또는 변수의 기능 설명 입력 후 저장
def funcdata_object_description(driver, object):
    if object == "함수":
        # 함수 정보 페이지 - 함수 탭 클릭 후 테이블 최상단 함수 클릭
        move_menu.safe_click(driver, By.ID, "tabFunc")
        driver.find_element(By.XPATH, '//*[@id="scmCodeTraceStructFunction"]/tr[1]/td[10]/img').click()
    elif object == "변수":
        # 함수 정보 페이지 - 변수 탭 클릭 후 테이블 최상단 변수 클릭
        driver.find_element(By.XPATH, '//div[6]/div[1]/div[2]/div/ul/li[2]/a').click()
        driver.find_element(By.XPATH, '//*[@id="scmCodeTraceStructValidation"]/tr[1]/td[9]/img').click()

    # 기능, 설명 입력 후 저장
    driver.find_element(By.ID, 'textareaFunction').send_keys(f'{object} 기능')
    driver.find_element(By.ID, 'textareaFunctionReason').send_keys(f'{object} 설명')

    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailScmCodeTraceExceptionModal"]/div/div/div[3]/a').click()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailScmCodeTraceExceptionModal"]/div/div/div[3]/button').click()
    time.sleep(1)


def upload_trial_exam_file(driver, file_category, file_name):
    wait_exam = WebDriverWait(driver, 15)
    if file_category == "src":
        driver.find_element(By.XPATH,'//*[@id="ProjectDetailSoftwareProcess-HeaderVue"]//li[1]/span[contains(text(),"체크섬")]').click()
        directory = os.getcwd()
        src_path = os.path.join(directory, 'trial_exam', 'src', file_name)
        driver.find_element(By.ID, 'uploadFile').send_keys(src_path)
        if len((driver.find_elements(By.XPATH, '//*[@id="step1"]/div[1]/div/div[2]/div/div/div/label[contains(text(),"파일을 선택해주세요.")]')))==0:
            time.sleep(1)
            driver.find_element(By.XPATH, '//*[@id="step1"]//button[contains(text(),"시험 결과 반영")]').click()
            time.sleep(5)
    elif file_category == "sps":
        driver.find_element(By.XPATH,'//*[@id="ProjectDetailSoftwareProcess-HeaderVue"]//li[1]/span[contains(text(),"체크섬")]').click()
        directory = os.getcwd()
        SPS_path = os.path.join(directory, 'trial_exam', 'SPS', file_name)
        driver.find_element(By.ID, 'spsFile').send_keys(SPS_path)
        if len((driver.find_elements(By.XPATH, '//*[@id="step1"]/div[1]/div/div[1]/div/div/div/label[contains(text(),"파일을 선택해주세요.")]')))==0:
            time.sleep(1)
            driver.find_element(By.XPATH, '//*[@id="step1"]//button[contains(text(),"시험 결과 반영")]').click()
            time.sleep(5)
    elif file_category == "static":
        driver.find_element(By.XPATH, '//*[@id="ProjectDetailSoftwareProcess-HeaderVue"]//li[2]/span[contains(text(),"정적")]').click()
        directory = os.getcwd()
        STATIC_xml_path = os.path.join(directory, 'xml', 'STATIC', file_name)
        driver.find_element(By.ID, 'staticXMLFile').send_keys(STATIC_xml_path)
        time.sleep(1)
        driver.find_element(By.XPATH,'//*[@id="step2"]//button[contains(text(),"시험 결과 반영")]').click()
        time.sleep(5)
    elif file_category == "dynamic":
        driver.find_element(By.XPATH, '//*[@id="ProjectDetailSoftwareProcess-HeaderVue"]//li[3]/span[contains(text(),"동적")]').click()
        directory = os.getcwd()
        CT_xml_path = os.path.join(directory, 'xml', 'CT', file_name)
        driver.find_element(By.ID, 'dynamicXMLFile').send_keys(CT_xml_path)
        time.sleep(1)
        driver.find_element(By.XPATH,'//*[@id="step3"]//button[contains(text(),"시험 결과 반영")]').click()
        time.sleep(5)


# 파일 정보 페이지 - 앞서 선택한 파일을 제외 처리
def filedata_select_change_except(driver, text):
    driver.find_element(By.XPATH, '//*[@id="projectInfoTable_wrapper"]/div[1]/button[2]').click()
    driver.find_element(By.XPATH, '//*[@id="projectInfoTable_wrapper"]/div[1]/div[2]/div/button[2]').click()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]/div/div/div[2]/form[2]/div/div/div[1]').click()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]//span[contains(text(),"제외")]').click()
    time.sleep(1)

    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]//textarea').clear()
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]//textarea').send_keys(text)
    driver.find_element(By.XPATH, '//*[@id="ProjectDetailFileManageTargetConfigModal"]/div/div/div[3]/button[2]').click()

# VPES 생성된 산출물 파일 다운로드 받을 때 [안전하지 않은 다운로드가 차단됨] 해제
def enable_download(self):
    download_dir = os.path.join(os.getcwd(), "report")

    chrome_options = Options()
    chrome_options.add_experimental_option("prefs", {
        "download.prompt_for_download": False,  # 다운로드 확인 팝업 비활성화
        "download.directory_upgrade": True,  # 기존 폴더 업그레이드 허용
        "safebrowsing.enabled": True,  # 안전 브라우징 활성화
        "safebrowsing.disable_download_protection": True  # 안전하지 않은 다운로드 허용
    })

    self.driver.execute_cdp_cmd("Page.enable", {})
    self.driver.execute_cdp_cmd("Browser.setDownloadBehavior", {
        "behavior": "allow",
        "downloadPath": download_dir
    })

# 실험실beta -> CRC32 hex 표기방식 사용
def convert_to_crc32Hex(driver, case_id):
    element = driver.find_element(By.XPATH, '//*[@id="TopNavBarAboutRtms"]/div')
    driver.execute_script("arguments[0].click();", element)  # 우측 상단 도움말 아이콘 클릭
    time.sleep(1)

    driver.find_element(By.XPATH, '//a[contains(text(),"실험실")]').click()  # [실험실beta] 클릭
    time.sleep(1)

    driver.find_element(By.ID, 'labBtnPrjChange').click()  # [변경] 드롭다운 클릭
    time.sleep(1)
    driver.find_element(By.XPATH, f'//*[@id="LaboratoryUiControlVue"]//a[contains(text(), "{case_id}")]').click()  # 해당 프로젝트 클릭
    time.sleep(1)

    # 코드 입력 : "LABCHECKSUM" 입력 후 활성화 클릭
    input_code = driver.find_element(By.ID, 'labInputActivateCode')
    input_code.send_keys("LABCHECKSUM")
    time.sleep(1)

    driver.find_element(By.ID, 'labBtnDoActivate').click()  # [활성화] 버튼 클릭
    time.sleep(1)

    driver.find_element(By.ID, 'labCheckBoxLabCommonChecksumDispalyHex').click()  # [CRC32 hex 표기방식 사용] 체크박스 클릭
    time.sleep(1)
    driver.find_element(By.ID, 'labCheckBoxLabCommonChecksumUseMD5').click()  # [MD5 Checksum 대체] 체크박스 클릭
    time.sleep(1)

def is_current_group_page(driver):
    current_url = driver.current_url
    return current_url.rstrip("/").split("/")[-1] == "ProjectGroup"

# 그룹 > 산출물 > 신뢰성 문서
def click_group_reliabilityReport(driver):
    if not is_current_group_page(driver):
        move_menu.move_project_group(driver)
    driver.find_element(By.ID, 'tabSN').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[1]').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[1]/div/div[3]/ul/li[1]').click()
    time.sleep(1)

# 그룹 > 산출물 > SW 문서
def click_group_swReport(driver):
    if not is_current_group_page(driver):
        move_menu.move_project_group(driver)
    driver.find_element(By.ID, 'tabSN').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[1]').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[1]/div/div[3]/ul/li[2]').click()
    time.sleep(1)

# 그룹 > 산출물 > KOLAS 문서
def click_group_kolasReport(driver):
    if not is_current_group_page(driver):
        move_menu.move_project_group(driver)
    driver.find_element(By.ID, 'tabSN').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[1]').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[1]/div/div[3]/ul/li[3]').click()
    time.sleep(1)

# 그룹 > 산출물 > 한글 개별
def click_group_koreanIndividual(driver):
    if not is_current_group_page(driver):
        move_menu.move_project_group(driver)
    driver.find_element(By.ID, 'tabSN').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]/div/div[3]/ul/li[1]').click()
    time.sleep(1)

# 그룹 > 산출물 > 한글 통합
def click_group_koreanIntegrated(driver):
    if not is_current_group_page(driver):
        move_menu.move_project_group(driver)
    driver.find_element(By.ID, 'tabSN').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]').click()
    driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]/div/div[3]/ul/li[2]').click()
    time.sleep(1)

# 그룹 > 산출물 > 워드
def click_group_word(driver, doc):
    if not is_current_group_page(driver):
        move_menu.move_project_group(driver)
    if doc == "신뢰성 문서":
        driver.find_element(By.ID, 'tabSN').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]/div/div[3]/ul/li[3]').click()
        time.sleep(1)
    elif doc == "SW 문서":
        driver.find_element(By.ID, 'tabSN').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]/div/div[3]/ul/li[2]').click()
        time.sleep(1)

# 그룹 > 산출물 > 엑셀
def click_group_excel(driver, doc):
    if not is_current_group_page(driver):
        move_menu.move_project_group(driver)
    if doc == "신뢰성 문서":
        driver.find_element(By.ID, 'tabSN').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]/div/div[3]/ul/li[4]').click()
        time.sleep(1)
    elif doc == "SW 문서":
        driver.find_element(By.ID, 'tabSN').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]').click()
        driver.find_element(By.XPATH, '//*[@id="TabGroupReport"]/div/div[2]/div[2]/div/div[3]/ul/li[3]').click()
        time.sleep(1)



# 회차 종료
def change_episode(driver, case_id, episode):
    move_menu.move_progress(driver, case_id)
    driver.find_element(By.XPATH, '//*[@id="projectDetailTransition-HeaderVue"]/div/div[2]/div[1]/div[2]/div[1]/div[2]').click()
    element = driver.find_element(By.XPATH, f"//*[@id='projectDetailTransition-HeaderVue']/div/div[2]/div/div[2]/div/div[3]//span/span[contains(text(),'{episode}')]")
    element.click()
    time.sleep(1.5)

# 프로젝트 설정 - CI & SN 빌드 설정
def project_build_setting(driver, case_id, type=None):

    # 프로젝트 설정 - 빌드 설정 탭 이동
    move_menu.move_buildSetting(driver, case_id)

    # 빌드 타입 - "Code Inspector & SNIPER" 선택
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[4]/div[2]/div/div[1]").click()
    driver.find_element(By.XPATH, "//*/text()[normalize-space(.)='Code Inspector & SNIPER']/parent::*").click()

    # CI 버전 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[1]/div/div[2]').click()
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[1]/div/div[3]/ul/li[1]/span/span").click()

    # 툴체인 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[2]/div/div[2]').click()
    driver.find_element(By.XPATH, "//li[8]/span/span").click()

    # 규칙 선택
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[3]/div/div[2]").click()
    time.sleep(0.2)
    driver.find_element(By.XPATH, "//div[3]/div/div[3]/ul/li[3]/span").click()
    time.sleep(0.2)

    # SN 버전 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[1]/div/div[1]').click()
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[1]/div/div[3]/ul/li[1]/span/span").click()

    # 툴체인 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[2]/div/div[2]').click()
    time.sleep(0.2)
    driver.find_element(By.XPATH, "//div[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div/div[2]/div[2]/div/div[3]/ul/li[8]/span/span").click()

    # 규칙 선택
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[3]/div/div[2]").click()
    driver.find_element(By.XPATH, "//div[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div/div[2]/div[3]/div/div[3]/ul/li/span/span").click()

    if type == "CI":
        driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[3]/div/img").click()  # SN 시험 제거
    elif type == "SN":
        driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[1]/div/img').click()  # CI 시험 제거

    # 빌드 설정 저장
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[1]/div/div[2]/div[2]/button").click()  # 저장

def project_build_setting_detail(driver, case_id, toolchain, CI_rule=None, SN_rule=None, type=None):
    if CI_rule == None:
        CI_rule = "CWE_C"
    elif SN_rule == None:
        SN_rule = "MISSION_CRITICAL_ERRORS"

    # 프로젝트 설정 - 빌드 설정 탭 이동
    move_menu.move_buildSetting(driver, case_id)

    # 빌드 타입 - "Code Inspector & SNIPER" 선택
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[4]/div[2]/div/div[1]").click()
    driver.find_element(By.XPATH, "//*/text()[normalize-space(.)='Code Inspector & SNIPER']/parent::*").click()

    # CI 버전 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[1]/div/div[2]').click()
    time.sleep(1)
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[1]/div/div[3]/ul/li[1]/span/span").click()

    # 툴체인 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[2]/div/div[2]').click()
    time.sleep(0.5)
    driver.find_element(By.XPATH, f"//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[2]/div/div[3]//span/span[contains(text(), '{toolchain}')]").click()

    # 규칙 선택
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[3]/div/div[2]").click()
    time.sleep(0.5)
    driver.find_element(By.XPATH, f"//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]/div[2]/div[3]/div/div[3]//span/span[contains(text(), '{CI_rule}')]").click()
    time.sleep(0.5)

    # SN 버전 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[1]/div/div[1]').click()
    time.sleep(0.5)
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[1]/div/div[3]/ul/li[1]/span/span").click()

    # 툴체인 선택
    driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[2]/div/div[2]').click()
    time.sleep(0.5)
    driver.find_element(By.XPATH, f"//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[2]/div/div[3]//span/span[contains(text(), '{toolchain}')]").click()

    # 규칙 선택
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[3]/div/div[2]").click()
    time.sleep(0.5)
    driver.find_element(By.XPATH, f"//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]/div[2]/div[3]/div/div[3]//span/span[contains(text(), '{SN_rule}')]").click()

    if type == "CI":
        driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[3]/div/img").click()  # SN 시험 제거
    elif type == "SN":
        driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[1]/div/img').click()  # CI 시험 제거

    # 빌드 설정 저장
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[1]/div/div[2]/div[2]/button").click()  # 저장
    time.sleep(1)

# 프로젝트 설정 - 방사청 개정안 선택
def project_doc_version_setting(driver, case_id, doc_version):
    if doc_version == "2022":
        version = "방사청 제2022-6호"
    elif doc_version == "2020":
        version = "방사청 제2020-1호"
    elif doc_version == "2018":
        version = "방사청 제2018-7호"

    # 프로젝트 설정 - 빌드 설정 탭 이동
    move_menu.move_projectAdditionalInfo(driver, case_id)

    driver.find_element(By.XPATH, '//*[@id="projectRegister-additionalOpt"]/div[8]/div/div[2]/div/div/div[1]').click()
    time.sleep(0.2)
    driver.find_element(By.XPATH, f"//*[@id='projectRegister-additionalOpt']/div[8]/div/div[2]/div/div/div[3]/ul//span[contains(text(), '{version}')]").click()

    driver.find_element(By.ID, "projectBtn").click()

# 설정 탭 > 프로젝트 설정 > 옵션 > autoCRFL [사용]
def use_autoCRLF(driver, case_id):
    move_menu.move_projectSetting(driver, case_id)  # 설정 탭 > 프로젝트 설정 클릭

    driver.find_element(By.ID, 'projectRegister-options-tab').click()  # [옵션] 클릭
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="projectRegister-options"]/div[1]/div[2]/div[2]/div/div/div[1]').click()
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="projectRegister-options"]/div[1]/div[2]/div[2]/div/div/div[3]/ul/li[1]').click()  # autoCRFL [사용] 클릭
    time.sleep(1)
    driver.find_element(By.ID, 'projectBtn').click()  # [저장] 버튼 클릭
    time.sleep(1)

# 설정 탭 > 프로젝트 설정 > 옵션 > autoCRFL [사용 안함]
def unuse_autoCRLF(driver, case_id):
    move_menu.move_projectSetting(driver, case_id)  # 설정 탭 > 프로젝트 설정 클릭

    driver.find_element(By.ID, 'projectRegister-options-tab').click()  # [옵션] 클릭
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="projectRegister-options"]/div[1]/div[2]/div[2]/div/div/div[1]').click()
    time.sleep(1)
    driver.find_element(By.XPATH, '//*[@id="projectRegister-options"]/div[1]/div[2]/div[2]/div/div/div[3]/ul/li[2]').click()  # autoCRFL [사용 안함] 클릭
    time.sleep(1)
    driver.find_element(By.ID, 'projectBtn').click()  # [저장] 버튼 클릭
    time.sleep(1)


def project_solutionbuild_setting_combined(driver, case_id, toolchain, Solution_Project, CI_rule, SN_rule, type, protype=None):
    """
    통합 빌드 설정 함수
    :param driver: Selenium WebDriver
    :param case_id: 케이스 식별자
    :param toolchain: 툴체인 이름
    :param Solution_Project: 솔루션 파일 이름 (예: "\\vs2015_miro.sln")
    :param CI_rule: CI용 규칙 (기본: CWE_C)
    :param SN_rule: SN용 규칙 (기본: MISSION_CRITICAL_ERRORS)
    :param type: CI만 = "CI", SN만 = "SN", 둘 다 = "ALL"
    """
    tool_setting(driver, case_id, type)
    if CI_rule == None:
        CI_rule = "CWE_C"
    if SN_rule == None:
        SN_rule = "MISSION_CRITICAL_ERRORS"
    if protype:
            project_type =1 #소스 파일 프로젝트
    else:
        project_type =2 # 솔루션 프로젝트

    # 빌드 설정 탭 이동
    move_menu.move_buildSetting(driver, case_id)

    # 빌드 타입 선택
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[4]/div[2]/div/div[1]").click()
    driver.find_element(By.XPATH, "//*/text()[normalize-space(.)='Code Inspector & SNIPER']/parent::*").click()

    def configure_section(section_xpath_base, rule_text):
        # 버전 선택
        driver.find_element(By.XPATH, f'{section_xpath_base}/div[2]/div[1]/div').click()
        time.sleep(0.2)
        driver.find_element(By.XPATH, f"{section_xpath_base}/div[2]/div[1]/div/div[3]").click()
        time.sleep(0.5)

        # 툴체인 선택
        driver.find_element(By.XPATH, f'{section_xpath_base}/div[2]/div[2]/div').click()
        time.sleep(0.2)
        driver.find_element(By.XPATH, f"{section_xpath_base}/div[2]/div[2]/div/div[3]//span/span[contains(text(), '{toolchain}')]").click()
        time.sleep(0.5)

        # 규칙 선택
        driver.find_element(By.XPATH, f'{section_xpath_base}/div[2]/div[3]/div').click()
        time.sleep(0.2)
        driver.find_element(By.XPATH, f"{section_xpath_base}/div[2]/div[3]/div/div[3]//span/span[contains(text(), '{rule_text}')]").click()

    # CI 설정
    if type in ["ALL", "CI"]:
        configure_section(
            "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[1]",
            CI_rule
        )
        # 프로젝트 타입 선택 및 솔루션 검색
        driver.find_element(By.XPATH, f"//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[2]/div/div[2]/div[{project_type}]/label").click()
        if project_type == 2:
            driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[2]/div/div[2]/button").click()
            # 이미지 검색 완료 대기
            while urlparse(driver.find_element(
                    By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[2]/div/div[2]/button/img"
            ).get_attribute("src")).path != "/vpes/image/svg/icon-link-off-copy-2.svg":
                time.sleep(1)
            # 솔루션 파일 선택
            time.sleep(2)
            driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[2]/div/div[2]/div[3]/div/div[1]").click()
            driver.find_element(By.XPATH, f'//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[2]/div[2]/div/div[2]/div[3]/div/div[3]/ul//span[contains(text(), "{Solution_Project}")]').click()

    # SN 설정
    if type in ["ALL", "SN"]:
        configure_section(
            "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[1]",
            SN_rule
        )
        driver.find_element(By.XPATH, f"//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[2]/div/div[2]/div[{project_type}]/label").click()
        if project_type == 2:
            driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[2]/div/div[2]/button").click()
            # 이미지 검색 완료 대기
            while urlparse(driver.find_element(
                    By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[2]/div[2]/div/div[2]/button/img"
            ).get_attribute("src")).path != "/vpes/image/svg/icon-link-off-copy-2.svg":
                time.sleep(1)
            # 솔루션 파일 선택
            time.sleep(2)
            driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[4]/div[2]/div/div[2]/div[3]/div/div[1]").click()
            driver.find_element(By.XPATH, f'//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[4]/div[2]/div/div[2]/div[3]/div/div[3]/ul//span[contains(text(), "{Solution_Project}")]').click()

    # CI 또는 SN 제거
    if type != 'ALL':
        if type == "CI":
            driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[2]/div/div[5]/div/div/div[3]/div[3]/div/img").click()
        elif type == "SN":
            driver.find_element(By.XPATH, '//*[@id="toolSettingModalVue"]/div[2]/div/div[5]/div/div/div[3]/div[1]/div/img').click()

    # 저장
    driver.find_element(By.XPATH, "//*[@id='toolSettingModalVue']/div[1]/div/div[2]/div[2]/button").click()

def build_run(driver, case_id):
    # 빌드 실행
    move_menu.move_buildExecution(driver, case_id)
    driver.find_element(By.ID, "btnBuildStart").click()

    # 빌드 상태 확인 대기
    progress_status = driver.find_element(By.XPATH, '//*[@id="projectSideBar"]/div/span[1]')
    while True:
        if progress_status.text == "진행중":
            break

#LLM 사용 시 로딩 스피너가 사라질 때까지 대기한 후 페이지를 새로고침
def wait_for_spinner_and_refresh(driver):
    spinner_selector = "//*[local-name()='svg' and contains(@style, 'spinner-border')]"

    while True:
        try:
            # 스피너가 존재하는지 확인
            spinner = driver.find_element(By.XPATH, spinner_selector)

            # 스피너가 화면에 표시되어 있으면 1초 대기
            if spinner.is_displayed():
                print("스피너가 활성화 상태입니다. 1초 대기...")
                time.sleep(1)
            else:
                print("스피너가 비활성화되었습니다.")
                break
        except NoSuchElementException:
            # 스피너 요소를 찾을 수 없을 경우 (즉, 비활성화되었거나 DOM에서 제거됨)
            print("스피너가 비활성화되었거나 사라졌습니다.")
            break

    driver.refresh()